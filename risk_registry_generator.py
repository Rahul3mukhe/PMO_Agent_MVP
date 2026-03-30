from __future__ import annotations

import io
import json
import os
import re
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from llm_providers import generate_text
from schemas import DocumentArtifact, PMOState

# =============================================================================
# PALETTE  (consistent with doc_templates.py)
# =============================================================================
_NAVY  = RGBColor(0x00, 0x33, 0x66)
_BLUE  = RGBColor(0x00, 0x72, 0xBB)
_BODY  = RGBColor(0x1A, 0x1A, 0x2E)
_MID   = RGBColor(0x4A, 0x55, 0x68)
_MUTED = RGBColor(0x6B, 0x72, 0x80)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_GREEN = RGBColor(0x15, 0x80, 0x3D)
_AMBER = RGBColor(0xB4, 0x53, 0x09)
_RED   = RGBColor(0xB9, 0x1C, 0x1C)

_HDR_BG   = "003366"
_ROW_A    = "EEF4FB"
_ROW_B    = "FFFFFF"
_RED_BG   = "FDE8E8"
_AMBER_BG = "FEF3C7"
_GREEN_BG = "E6F4EA"
_CRIT_BG  = "FCCFCF"


# =============================================================================
# SECTION 1 — LLM PROMPT
# =============================================================================

def _build_prompt(state: PMOState) -> str:
    proj = state.project
    org  = state.standards["org"]["name"]

    # Collect ALL uploaded content — explicit mapping AND raw upload text
    uploaded_parts = []
    if "uploaded_mapping" in state.audit:
        mapped = state.audit["uploaded_mapping"].get("risk_registry", "")
        if mapped and mapped.strip():
            uploaded_parts.append(mapped.strip())
    raw_upload = state.audit.get("raw_upload_text", "")
    if raw_upload and raw_upload.strip():
        uploaded_parts.append(raw_upload.strip())
    uploaded = "\n\n".join(uploaded_parts)

    known = "\n".join(f"  - {r}" for r in (proj.known_risks or []))
    deliv = "\n".join(f"  - {d}" for d in (proj.key_deliverables or []))

    return f"""You are a senior PMO Risk Specialist at {org} with 20 years of experience
delivering complex {proj.project_type} projects. Your job is to produce a comprehensive,
realistic Risk Registry for the project below — exactly as a real risk specialist would
in a formal audit-ready PMO document.

IMPORTANT: Think deeply about the project type, industry, and scope. Identify ALL realistic
risks this specific project faces. For example, a banking digital transformation project
would face: core banking system stability risks, PCI-DSS and GDPR compliance risks,
real-time payment processing risks, data migration risks, regulatory approval risks,
legacy API deprecation risks, fraud risk exposure during transition, etc.

━━━ PROJECT CONTEXT ━━━
Organisation     : {org}
Project ID       : {proj.project_id}
Project Name     : {proj.project_name}
Project Type     : {proj.project_type}
Sponsor          : {proj.sponsor or "To be confirmed"}
Estimated Budget : {proj.estimated_budget or "Not specified"}
Actual Spend     : {proj.actual_budget_consumed or "Not yet available"}
Timeline (days)  : {proj.total_time_taken_days or "Ongoing"}
Timeline Summary : {proj.timeline_summary or "Not provided"}
Scope Summary    : {proj.scope_summary or "Not provided"}
Key Deliverables :
{deliv or "  (none listed)"}
Risks already known to the project team (MUST be included with source = Human Identified):
{known or "  (none — derive all risks from the project context above)"}

━━━ UPLOADED DOCUMENT CONTEXT ━━━
{uploaded[:5000] if uploaded else "(none — derive risks entirely from the project description above)"}

━━━ YOUR TASK ━━━
Read the project description carefully. Identify ALL realistic risks this project faces:
  - Technical risks (integration, performance, architecture, data)
  - Regulatory and compliance risks (GDPR, PCI-DSS, SOX, MAS TRM, or relevant regulations)
  - Financial risks (budget overrun, cost escalation, ROI)
  - Security risks (cyber, data breach, access control)
  - Resource risks (attrition, skill gaps, availability)
  - Schedule risks (dependency delays, critical path)
  - Vendor and third-party risks (SLA breach, dependency)
  - Organisational and change management risks (adoption, resistance)
  - Operational risks (business continuity, rollback)

OWNER RULE (CRITICAL):
  - Set "owner" to exactly "Project Manager" for EVERY risk — no exceptions.
  - Set "owner_note" to the role that should eventually own this risk
    e.g. "Reassign to: Head of Information Security" or "Reassign to: Finance Lead"

━━━ OUTPUT FORMAT ━━━
Return ONLY a single valid JSON object. No markdown, no explanation, no code fences.

{{
  "executive_summary": "3-4 sentences summarising the overall risk posture for {proj.project_name}. Be specific about the project, org, and top risk areas.",

  "risk_framework": "2 sentences: methodology is ISO 31000, 5x5 likelihood-impact matrix, risks escalate based on score thresholds, owners reassigned from PM after review.",

  "risk_posture": "ELEVATED | MODERATE | CONTROLLED",

  "assumptions": [
    "Specific assumption 1 relevant to this project",
    "Specific assumption 2",
    "At least 8 assumptions total"
  ],

  "risks": [
    {{
      "id": "R-001",
      "source": "AI Identified",
      "category": "Technical",
      "title": "Concise risk title max 10 words",
      "description": "2-3 sentences: what the risk is, root cause, and specific consequence for this project.",
      "trigger": "The specific event or condition that would cause this risk to materialise.",
      "likelihood": "Low | Medium | High",
      "likelihood_score": 3,
      "impact": "Low | Medium | High | Critical",
      "impact_score": 4,
      "rating": "High",
      "score": 12,
      "owner": "Project Manager",
      "owner_note": "Reassign to: [appropriate specialist role]",
      "status": "Open",
      "mitigation": "3-4 specific actionable steps referencing this project context.",
      "contingency": "2-3 specific actions if the risk materialises — who does what.",
      "assumption": "The key assumption this risk assessment depends on.",
      "review_frequency": "Weekly | Fortnightly | Monthly",
      "residual_rating": "Low | Medium | High"
    }}
  ],

  "review_cadence": [
    {{
      "review_type": "Review name",
      "frequency": "How often",
      "attendees": "Who attends",
      "scope": "What is reviewed",
      "output": "Document or decision produced"
    }}
  ]
}}

RULES:
1. Minimum 10 risks. For complex or regulated projects generate 12-15.
2. score = likelihood_score x impact_score (both 1-5). Rating: 1-3=Low, 4-8=Medium, 9-15=High, 16-25=Critical.
3. Human-identified risks use source "Human Identified". AI risks use "AI Identified".
4. owner is ALWAYS exactly "Project Manager". Never any other value.
5. owner_note tells the PM who should eventually own this risk.
6. Every single field must be populated. No nulls, no empty strings.
7. Output ONLY the JSON object. Nothing before or after.
""".strip()


# =============================================================================
# SECTION 2 — LLM CALL + PARSE + VALIDATE
# =============================================================================

def _parse_json(raw: str) -> Dict[str, Any]:
    text  = re.sub(r"```(?:json)?", "", raw).strip().rstrip("`").strip()
    start = text.find("{")
    end   = text.rfind("}") + 1
    if start == -1 or end == 0:
        raise ValueError("No JSON object in LLM response")
    return json.loads(text[start:end])


def _fill_defaults(data: Dict[str, Any], proj) -> Dict[str, Any]:
    data.setdefault("executive_summary",
                    f"Risk registry for {proj.project_name} ({proj.project_id}).")
    data.setdefault("risk_framework",
                    "Risks assessed using ISO 31000 with a 5x5 likelihood-impact matrix.")
    data.setdefault("risk_posture", "ELEVATED")
    data.setdefault("assumptions", [
        "Project scope is baselined and change-controlled after sign-off.",
        "Key named resources remain available for the full project duration.",
        "Regulatory requirements do not change materially during delivery.",
        "Steering committee convenes at scheduled intervals without cancellation.",
        "Budget contingency reserve is pre-approved and accessible upon PMO request.",
        "Third-party vendors meet contractual SLA obligations.",
        "Executive sponsorship remains active and visible throughout delivery.",
        "All project gates are reviewed within agreed timeframes.",
    ])
    data.setdefault("review_cadence", [
        {"review_type": "Weekly Risk Standup",
         "frequency": "Weekly", "attendees": "PM + Risk Owners",
         "scope": "Critical and High risks", "output": "Risk Action Log"},
        {"review_type": "Fortnightly Risk Review",
         "frequency": "Fortnightly", "attendees": "PM + Delivery Leads",
         "scope": "All open risks", "output": "Updated Risk Register"},
        {"review_type": "Steering Committee Escalation",
         "frequency": "Monthly", "attendees": "Sponsor + SteerCo + PM",
         "scope": "Critical escalations and decisions", "output": "SteerCo Minutes"},
        {"review_type": "Full Register Audit",
         "frequency": "Quarterly", "attendees": "PMO + Internal Audit",
         "scope": "Complete risk register", "output": "Audit Report"},
    ])

    risk_defaults = {
        "id": "R-000", "source": "AI Identified", "category": "General",
        "title": "Unnamed Risk", "description": "To be documented.",
        "trigger": "To be identified.", "likelihood": "Medium",
        "likelihood_score": 3, "impact": "Medium", "impact_score": 3,
        "rating": "Medium", "score": 9,
        "owner": "Project Manager",
        "owner_note": "Reassign to: appropriate specialist after PM review.",
        "status": "Open",
        "mitigation": "To be defined by risk owner.",
        "contingency": "To be defined by risk owner.",
        "assumption": "Standard project assumptions apply.",
        "review_frequency": "Monthly", "residual_rating": "Low",
    }
    for i, r in enumerate(data.get("risks", [])):
        for k, v in risk_defaults.items():
            if k not in r or r[k] in (None, ""):
                r[k] = v
        if r["id"] == "R-000":
            r["id"] = f"R-{str(i+1).zfill(3)}"
        # Enforce PM as owner regardless of what LLM returned
        r["owner"] = "Project Manager"

    return data



def _builtin_risks(proj) -> Dict[str, Any]:
    """
    Generate a realistic, project-specific risk register from the
    project fields alone — no LLM required.

    This is used when:
      • All LLM providers fail (Groq + Ollama + Mistral all down)
      • LLM returns non-JSON (local template markdown fallback fired)
      • LLM returns JSON but with 0 risks

    Output is structured identically to the LLM JSON schema so it
    flows through the same renderer unchanged.
    """
    name    = getattr(proj, "project_name",  "This Project")
    ptype   = getattr(proj, "project_type",  "Software Development")
    budget  = getattr(proj, "estimated_budget", None)
    actual  = getattr(proj, "actual_budget_consumed", None)
    days    = getattr(proj, "total_time_taken_days", None)
    scope   = getattr(proj, "scope_summary",  "") or ""
    tline   = getattr(proj, "timeline_summary", "") or ""
    deliv   = getattr(proj, "key_deliverables", []) or []
    known   = getattr(proj, "known_risks", []) or []
    pid     = getattr(proj, "project_id", "PRJ-000")
    sponsor = getattr(proj, "sponsor", "Project Sponsor") or "Project Sponsor"

    # Derive context clues from text fields
    scope_lower = (scope + " " + tline + " " + name + " " + ptype).lower()
    is_banking   = any(w in scope_lower for w in ["bank","payment","transaction","pci","fca","basel","lending","credit","fintech","financial service"])
    is_cloud     = any(w in scope_lower for w in ["aws","azure","gcp","cloud","kubernetes","docker","microservice","serverless"])
    is_mobile    = any(w in scope_lower for w in ["mobile","ios","android","react native","flutter","app"])
    is_regulated = any(w in scope_lower for w in ["gdpr","pci","fca","mas trm","sox","hipaa","iso 27001","regulated","compliance","regulatory"])
    is_migration = any(w in scope_lower for w in ["migrat","legacy","mainframe","cobol","cutover","transition","uplift"])
    is_api       = any(w in scope_lower for w in ["api","integration","third.party","vendor","partner","webhook"])
    is_data      = any(w in scope_lower for w in ["data","database","record","etl","pipeline","warehouse"])

    budget_str  = f"${budget:,.0f}" if budget else "the approved budget"
    actual_str  = f"${actual:,.0f}" if actual else "current spend"
    days_str    = f"{days} days" if days else "the planned timeline"
    deliv_str   = deliv[0] if deliv else f"{name} deliverables"

    # ── Build risk list ──────────────────────────────────────────────
    risks = []
    ri = [0]  # mutable counter for closure

    def add(source, category, title, desc, trigger, lh, ls, imp, ims,
            rating, score, mitigation, contingency, assumption,
            freq, residual, note):
        ri[0] += 1
        risks.append({
            "id": f"R-{str(ri[0]).zfill(3)}",
            "source": source,
            "category": category,
            "title": title,
            "description": desc,
            "trigger": trigger,
            "likelihood": lh, "likelihood_score": ls,
            "impact": imp,    "impact_score": ims,
            "rating": rating, "score": score,
            "owner": "Project Manager",
            "owner_note": note,
            "status": "Open",
            "mitigation": mitigation,
            "contingency": contingency,
            "assumption": assumption,
            "review_frequency": freq,
            "residual_rating": residual,
        })

    # ── 1. Human-identified risks first ─────────────────────────────
    for kr in known:
        add("Human Identified", "General",
            kr[:80] if len(kr) > 80 else kr,
            f"This risk was formally identified by the project team for {name}: {kr}. "
            f"It requires active monitoring, a documented mitigation plan, and assignment "
            f"to the appropriate specialist for ongoing management.",
            f"The conditions described in the risk statement materialise during {name} delivery.",
            "Medium", 3, "High", 4, "High", 12,
            f"1. Formally assess and document the full impact of this risk on {name}. "
            f"2. Assign to the appropriate domain specialist within 5 business days. "
            f"3. Define and implement specific mitigation actions with tracked deadlines. "
            f"4. Review status at every fortnightly risk review until the risk is closed.",
            f"1. Escalate immediately to {sponsor} and the PMO. "
            f"2. Convene an emergency risk review within 48 hours. "
            f"3. Activate the relevant contingency plan and update the RAID log.",
            f"The project team has provided sufficient context to assess and manage this risk.",
            "Fortnightly", "Low",
            "Reassign to: appropriate domain specialist after PM review")

    # ── 2. Schedule risk (almost always applicable) ──────────────────
    add("AI Identified", "Schedule",
        f"Schedule Delay Impacting {deliv_str[:50]}",
        f"{name} is currently tracking against {days_str}. Schedule delays — driven "
        f"by dependency failures, resource gaps, or scope changes — risk missing the "
        f"committed delivery milestones and triggering penalty clauses or stakeholder "
        f"confidence loss.",
        f"A key milestone is missed by more than 5 business days, pushing the critical path "
        f"and compressing downstream phases.",
        "High", 4, "High", 4, "High", 16,
        f"1. Maintain a weekly critical path review with all delivery leads. "
        f"2. Identify and formally document all schedule dependencies at the start of each phase. "
        f"3. Maintain a 10% schedule buffer on the critical path. "
        f"4. Escalate any delay exceeding 3 days to {sponsor} immediately.",
        f"1. Convene an emergency re-planning session within 48 hours of confirmed delay. "
        f"2. Descope non-critical features to protect the core delivery commitment. "
        f"3. Brief {sponsor} and stakeholders within 24 hours.",
        "All external dependencies will be delivered within their agreed SLA windows.",
        "Weekly", "Medium",
        "Reassign to: Delivery Lead")

    # ── 3. Budget overrun ────────────────────────────────────────────
    if budget:
        pct = round((actual / budget * 100), 1) if actual and budget else 0
        add("AI Identified", "Financial",
            f"Budget Overrun on {name}",
            f"The approved budget for {name} is {budget_str}, with {actual_str} consumed "
            f"({pct}% utilisation). Unplanned scope additions, resource cost escalation, "
            f"or vendor overruns risk breaching the budget envelope and requiring emergency "
            f"Investment Committee approval.",
            f"Monthly spend tracking shows actual costs exceeding the approved monthly baseline "
            f"by more than 10% for two consecutive months.",
            "Medium", 3, "High", 4, "High", 12,
            f"1. Implement monthly budget reviews as a standing agenda item at the Programme Board. "
            f"2. Enforce strict change control — all scope additions require a formal budget impact assessment. "
            f"3. Maintain a pre-approved contingency reserve of 10% of {budget_str}. "
            f"4. Configure automated alerts when actual spend exceeds 80% of any monthly budget allocation.",
            f"1. Release the pre-approved contingency reserve with {sponsor} notification. "
            f"2. Initiate an emergency budget review and descope non-critical work. "
            f"3. Escalate to the Investment Committee if the overrun exceeds 15% of {budget_str}.",
            f"A pre-approved contingency reserve of 10% of {budget_str} is accessible upon PMO request.",
            "Fortnightly", "Low",
            "Reassign to: Finance Lead")

    # ── 4. Resource / key person risk ────────────────────────────────
    add("AI Identified", "Resource",
        f"Key Resource Attrition or Unavailability",
        f"The delivery of {name} depends on a small number of specialists with domain "
        f"knowledge that is not fully documented or transferable at short notice. "
        f"Unplanned departure or unavailability of a key team member would create a "
        f"delivery gap that could delay milestones by 4-8 weeks.",
        f"A named key resource provides notice, goes on extended leave, or becomes "
        f"unavailable during a critical delivery phase.",
        "Medium", 3, "High", 4, "High", 12,
        f"1. Identify the top 3 key-person dependencies and document their knowledge in a "
        f"shared knowledge base within 30 days. "
        f"2. Assign at least one shadow or backup resource to each critical role. "
        f"3. Include knowledge transfer milestones in the project plan. "
        f"4. Confirm bench resource availability with the SI partner before Phase 2.",
        f"1. Activate the SI partner bench resource within 2 weeks. "
        f"2. Re-plan the affected workstream with the remaining team. "
        f"3. Brief {sponsor} within 24 hours of confirmed departure.",
        "A qualified bench resource is available from the SI partner within 2 weeks of a request.",
        "Monthly", "Low",
        "Reassign to: Delivery Manager")

    # ── 5. Third-party / vendor risk ─────────────────────────────────
    if is_api or len(deliv) > 1:
        add("AI Identified", "Vendor",
            f"Third-Party Vendor or API Delivery Delay",
            f"{name} depends on one or more third-party vendors or external APIs for key "
            f"deliverables. Vendor delays, API certification failures, or SLA breaches "
            f"would directly impact the project critical path and go-live readiness.",
            f"A third-party vendor misses a contracted delivery milestone or delivers an "
            f"API or component with defects requiring remediation.",
            "Medium", 3, "High", 4, "High", 12,
            f"1. Insert contractual penalty clauses and SLA milestones for all third-party "
            f"deliveries — confirm in final contracts before Phase 2. "
            f"2. Establish weekly delivery cadence calls with each vendor. "
            f"3. Commission a parallel evaluation of at least one alternative vendor for "
            f"any single-source dependency. "
            f"4. Request access to vendor sandbox environments in Phase 1 to begin early integration.",
            f"1. Activate the alternative vendor within 10 business days of confirmed delay. "
            f"2. Build internal stubs for vendor components to maintain UAT progress. "
            f"3. Negotiate schedule relief with {sponsor} if the critical path is impacted.",
            "All third-party vendors will meet contractual SLA obligations on time.",
            "Weekly", "Medium",
            "Reassign to: Delivery Lead")

    # ── 6. Security risk ─────────────────────────────────────────────
    add("AI Identified", "Security",
        f"Security Vulnerability or Data Breach in {name}",
        f"The {name} platform will handle sensitive data. Without adequate security "
        f"controls, penetration testing, and access management, the system is exposed "
        f"to data breaches, unauthorised access, and non-compliance with applicable "
        f"security standards (ISO 27001, OWASP, PCI-DSS, GDPR as applicable).",
        f"A security vulnerability is identified in penetration testing or a data breach "
        f"is detected in the live or staging environment.",
        "Low", 2, "Critical", 5, "High", 10,
        f"1. Conduct a full penetration test before go-live — mandatory gate deliverable. "
        f"2. Implement role-based access control (RBAC) across all system components. "
        f"3. Conduct OWASP Top 10 security review during development. "
        f"4. Deploy a SIEM solution covering all production components before go-live. "
        f"5. Establish a vulnerability management process with SLA-bound remediation.",
        f"1. Isolate the affected component immediately upon breach detection. "
        f"2. Activate the incident response plan within 1 hour. "
        f"3. Notify the relevant regulatory authority within required timeframes. "
        f"4. Engage external cyber specialists within 4 hours.",
        "The security penetration test will be completed and all critical findings remediated before go-live.",
        "Monthly", "Low",
        "Reassign to: Head of Information Security")

    # ── 7. Scope creep ───────────────────────────────────────────────
    add("AI Identified", "Organisational",
        f"Uncontrolled Scope Creep on {name}",
        f"Stakeholder requests for additional features or changes to the {name} scope "
        f"after baseline sign-off risk expanding the delivery commitment without a "
        f"corresponding increase in budget or timeline, eroding quality and increasing "
        f"team pressure.",
        f"A stakeholder or business unit requests a new feature or change that falls "
        f"outside the baselined scope without submitting a formal change request.",
        "High", 4, "Medium", 3, "High", 12,
        f"1. Establish a formal change control board (CCB) before Phase 2 commences. "
        f"2. Require all scope changes to be submitted via a Change Request form with "
        f"a mandatory budget and schedule impact assessment. "
        f"3. Brief {sponsor} on the cost of scope changes at each Programme Board. "
        f"4. Publish the baselined scope to all stakeholders and obtain written sign-off.",
        f"1. Decline the scope change formally in writing until a Change Request is raised. "
        f"2. Escalate to {sponsor} if the requestor bypasses the change control process. "
        f"3. Record all refused changes in the project log for audit purposes.",
        "All scope changes will be formally assessed and approved before being added to the delivery backlog.",
        "Fortnightly", "Low",
        "Reassign to: PMO Representative")

    # ── 8. Quality / testing risk ────────────────────────────────────
    add("AI Identified", "Technical",
        f"Insufficient Testing Coverage Leading to Production Defects",
        f"Inadequate test coverage — particularly for integration, performance, and "
        f"regression scenarios — risks releasing {name} with critical defects. "
        f"Post-go-live defects in production would damage user confidence, require "
        f"emergency patches, and increase support costs.",
        f"UAT sign-off is granted without full test coverage, or a critical defect is "
        f"identified in production within 30 days of go-live.",
        "Medium", 3, "High", 4, "High", 12,
        f"1. Define a formal test strategy with minimum coverage targets before Phase 2. "
        f"2. Implement automated regression testing covering all critical user journeys. "
        f"3. Conduct load and performance testing at 150% of expected peak capacity. "
        f"4. Enforce a zero-severity-1 defect threshold for go-live approval. "
        f"5. Extend UAT to a minimum of 4 weeks with named business sign-off.",
        f"1. Halt the go-live if any Severity 1 defects are open at the go/no-go review. "
        f"2. Activate the defect triage war room within 2 hours of a production incident. "
        f"3. Prepare a hotfix release within 48 hours of a critical production defect.",
        "All Severity 1 and Severity 2 defects will be resolved before UAT sign-off.",
        "Weekly", "Low",
        "Reassign to: QA Lead")

    # ── 9. Cloud-specific if applicable ─────────────────────────────
    if is_cloud:
        add("AI Identified", "Technical",
            f"Cloud Infrastructure Cost Overrun and Performance Issues",
            f"Cloud infrastructure costs for {name} are difficult to forecast before "
            f"production workloads are profiled. Unoptimised configurations, higher data "
            f"egress, and over-provisioned environments commonly track 20-35% above initial "
            f"estimates, compounding the financial risk.",
            f"Monthly cloud cost report shows actual spend exceeding the approved monthly "
            f"baseline by more than 15% for two consecutive months.",
            "High", 4, "Medium", 3, "High", 12,
            f"1. Engage a FinOps specialist before Phase 2 to baseline cloud costs. "
            f"2. Purchase reserved instances for all production workloads. "
            f"3. Configure automated cost anomaly alerts at 80% and 100% of monthly baseline. "
            f"4. Enforce automated shutdown of dev/test environments outside business hours.",
            f"1. Release the pre-approved contingency reserve. "
            f"2. Conduct an emergency FinOps review within 5 business days. "
            f"3. Right-size all non-production environments immediately.",
            "Reserved instances reduce cloud costs by at least 30% versus on-demand pricing.",
            "Fortnightly", "Low",
            "Reassign to: Finance Lead / Cloud Engineering Lead")

    # ── 10. Mobile-specific if applicable ───────────────────────────
    if is_mobile:
        add("AI Identified", "Technical",
            f"App Store Rejection or Delayed Platform Certification",
            f"The {name} mobile application requires Apple App Store and Google Play Store "
            f"certification before release. Rejection due to policy violations, security "
            f"issues, or metadata problems can delay go-live by 2-6 weeks per rejection cycle.",
            f"Apple or Google returns a rejection notice citing policy violations, security "
            f"vulnerabilities, or incomplete metadata during app submission.",
            "Medium", 3, "High", 4, "High", 12,
            f"1. Conduct a pre-submission review against Apple App Store Review Guidelines "
            f"and Google Play Policy before the first submission attempt. "
            f"2. Ensure all biometric authentication flows comply with platform-specific "
            f"security APIs (Face ID, Touch ID, Android Biometric). "
            f"3. Build a 4-week buffer into the go-live plan for app store review cycles. "
            f"4. Submit for TestFlight/Google Play internal testing at least 8 weeks before go-live.",
            f"1. Address all rejection reasons within 5 business days and resubmit. "
            f"2. Engage Apple/Google developer support for expedited review if timeline is critical. "
            f"3. Notify {sponsor} and adjust go-live date if rejection resolution exceeds 2 weeks.",
            "The first app store submission will be made at least 6 weeks before the planned go-live date.",
            "Weekly", "Low",
            "Reassign to: Lead Mobile Developer")

    # ── 11. Regulatory if applicable ────────────────────────────────
    if is_regulated or is_banking:
        add("AI Identified", "Regulatory",
            f"Regulatory Compliance Gap Identified Pre or Post Go-Live",
            f"{name} operates in a regulated environment. Non-compliance with applicable "
            f"regulations (GDPR, PCI-DSS, FCA requirements, or relevant standards) risks "
            f"enforcement action, financial penalties, mandatory remediation, and reputational damage.",
            f"Internal audit or a regulatory authority identifies a compliance gap in the "
            f"{name} system or processes before or after go-live.",
            "Medium", 3, "Critical", 5, "High", 15,
            f"1. Conduct a formal compliance impact assessment during Phase 1 covering all "
            f"applicable regulations for {name}. "
            f"2. Include regulatory compliance sign-off as a mandatory Phase 2 gate deliverable. "
            f"3. Engage external regulatory counsel for a pre-go-live compliance review. "
            f"4. Implement automated compliance monitoring in the production environment.",
            f"1. Halt any non-compliant processing immediately. "
            f"2. Engage external regulatory counsel within 24 hours. "
            f"3. Submit a voluntary disclosure to the relevant authority if a breach is confirmed. "
            f"4. Define and deliver a remediation plan within 10 business days.",
            "All applicable regulatory requirements are known and documented before Phase 2 commences.",
            "Fortnightly", "Low",
            "Reassign to: Chief Compliance Officer")

    # ── 12. Data migration if applicable ────────────────────────────
    if is_migration or is_data:
        add("AI Identified", "Technical",
            f"Data Migration Failure or Data Integrity Loss",
            f"Data migration activities for {name} carry a risk of data loss, corruption, "
            f"or reconciliation failures if the migration process is not thoroughly tested "
            f"and validated. Post-migration data integrity issues would undermine the "
            f"credibility of the new system and may require costly rollback.",
            f"Post-migration reconciliation checks identify a mismatch between source and "
            f"target datasets, or data integrity validation failures exceed the defined tolerance.",
            "Medium", 3, "Critical", 5, "High", 15,
            f"1. Conduct a minimum of 3 full migration rehearsals in the staging environment "
            f"before live migration — each must achieve 100% reconciliation. "
            f"2. Define a zero-tolerance threshold for data integrity errors. "
            f"3. Retain a point-in-time snapshot of all source data before live migration. "
            f"4. Assign a dedicated data quality lead accountable for migration integrity.",
            f"1. Halt migration immediately if the error threshold is breached. "
            f"2. Restore from the pre-migration snapshot — target RTO of 4 hours. "
            f"3. Conduct root cause analysis before retrying migration.",
            "The source system can produce a consistent point-in-time snapshot within the migration window.",
            "Weekly", "Low",
            "Reassign to: Head of Data Engineering")

    # ── Ensure we have at least 8 risks (add generic if very short) ──
    if len(risks) < 8:
        add("AI Identified", "Organisational",
            "Stakeholder Alignment and Governance Failure",
            f"Insufficient stakeholder engagement or governance failures during {name} "
            f"can result in delayed approvals, conflicting requirements, and lack of "
            f"executive support, undermining delivery confidence and gate progression.",
            "A key stakeholder withholds approval at a project gate or escalates conflicting requirements.",
            "Medium", 3, "Medium", 3, "Medium", 9,
            f"1. Define a formal RACI matrix and governance structure before Phase 2. "
            f"2. Schedule monthly stakeholder briefings with {sponsor}. "
            f"3. Document all stakeholder decisions in the project log. "
            f"4. Ensure all gate reviews have quorum and defined decision criteria.",
            f"1. Escalate to {sponsor} within 24 hours of a governance failure. "
            f"2. Convene an emergency stakeholder workshop within 5 business days. "
            f"3. Escalate to the Board if {sponsor}-level resolution is not achieved.",
            "All key stakeholders have been identified and their availability confirmed.",
            "Monthly", "Low",
            "Reassign to: PMO Representative")

    # ── Assumptions ──────────────────────────────────────────────────
    assumptions = [
        f"The project scope for {name} is fully baselined and approved before Phase 2 commences.",
        f"All key resources for {name} remain available throughout the delivery lifecycle.",
        f"The approved budget of {budget_str} includes a pre-approved contingency reserve of 10%." if budget else "A pre-approved contingency reserve is accessible upon PMO approval.",
        f"Third-party vendors and partners will meet their contractual SLA obligations.",
        f"Steering committee meetings will convene as scheduled without cancellation.",
        f"Regulatory and compliance requirements applicable to {name} will not change materially during delivery.",
        f"Executive sponsorship from {sponsor} will remain active and visible throughout the programme.",
        f"All project gate reviews will be completed within 5 business days of the scheduled date.",
        f"The delivery timeline of {days_str} is achievable with the current team size and composition." if days else "The planned delivery timeline is achievable with the current team and budget.",
        f"UAT business representatives will be available for a minimum of 20% time allocation during the test phase.",
    ]

    # ── Review cadence ───────────────────────────────────────────────
    review_cadence = [
        {"review_type": f"{name} Weekly Risk Standup",
         "frequency": "Weekly — every Monday",
         "attendees": f"Project Manager, Risk Owners, Delivery Leads",
         "scope": "All Critical and High risks — status, actions, new risks",
         "output": "Risk Action Log (updated in project RAID log)"},
        {"review_type": "Fortnightly Risk Register Review",
         "frequency": "Fortnightly — aligned to Sprint/Phase reviews",
         "attendees": f"Project Manager, All Risk Owners, PMO Representative",
         "scope": "Full register — rating updates, new risks, closed risks",
         "output": f"Updated Risk Register v{{n}} — circulated to {sponsor}"},
        {"review_type": "Monthly Steering Committee Escalation",
         "frequency": "Monthly",
         "attendees": f"{sponsor}, Steering Committee, Project Manager",
         "scope": "Critical and High escalations, investment decisions",
         "output": "SteerCo Minutes and Approved Actions"},
        {"review_type": "Quarterly Full Register Audit",
         "frequency": "Quarterly",
         "attendees": "PMO Director, Internal Audit, Project Manager",
         "scope": "Complete register — methodology, residuals, owner confirmation",
         "output": "Quarterly Risk Audit Report"},
    ]

    # ── Executive summary ─────────────────────────────────────────────
    crit  = sum(1 for r in risks if r["rating"] == "Critical")
    high  = sum(1 for r in risks if r["rating"] == "High")
    med   = sum(1 for r in risks if r["rating"] == "Medium")
    human = sum(1 for r in risks if r["source"] == "Human Identified")
    ai_n  = len(risks) - human

    posture = "ELEVATED" if (crit > 0 or high >= 3) else ("MODERATE" if high > 0 else "CONTROLLED")

    exec_summary = (
        f"{name} ({pid}) carries a {posture} risk posture across its delivery lifecycle. "
        f"A total of {len(risks)} risks have been identified: {human} raised directly by "
        f"the project team and {ai_n} identified through structured PMO risk analysis of "
        f"the project scope, deliverables, and operating environment. "
    )
    if crit > 0 or high > 0:
        exec_summary += (
            f"Of these, {crit} risk(s) are rated Critical and {high} are rated High, "
            f"requiring immediate attention from the Project Manager and {sponsor}. "
        )
    exec_summary += (
        f"All {len(risks)} risks are currently owned by the Project Manager pending "
        f"formal reassignment to the appropriate domain specialists. "
        f"The Project Manager must complete risk ownership reassignment within "
        f"5 business days of this document being approved."
    )

    return {
        "executive_summary": exec_summary,
        "risk_framework": (
            "This register is maintained in accordance with ISO 31000:2018 Risk Management "
            "Guidelines. Risks are scored on a 5x5 likelihood-impact matrix; scores of "
            "16-25 are Critical, 9-15 High, 4-8 Medium, and 1-3 Low. All risks are "
            "initially owned by the Project Manager pending formal reassignment."
        ),
        "risk_posture": posture,
        "assumptions": assumptions,
        "risks": risks,
        "review_cadence": review_cadence,
    }

def _call_llm(state: PMOState) -> Dict[str, Any]:
    prompt = _build_prompt(state)
    try:
        raw = generate_text(
            provider=state.provider,
            model=state.model,
            prompt=prompt,
            project=state.project,
            doc_type="risk_registry",
            standards=state.standards,
            api_key=state.audit.get("api_key"),
            temperature=0.2,
            max_tokens=state.audit.get("max_tokens", 4096),
        )
    except Exception as e:
        state.audit["risk_registry_llm_call_error"] = str(e)
        raw = ""

    # Detect if we got back the local template markdown (no JSON object) —
    # this happens when all real LLM providers fail and the local template
    # fallback fires. In that case go straight to built-in risk generation.
    has_json = raw.strip() and "{" in raw and "}" in raw
    if not has_json:
        state.audit["risk_registry_used_builtin"] = True
        return _fill_defaults(_builtin_risks(state.project), state.project)

    try:
        data = _parse_json(raw)
        # Sanity check — if LLM returned JSON but with 0 risks, use built-in
        if not data.get("risks"):
            state.audit["risk_registry_used_builtin"] = "LLM returned 0 risks"
            data = _builtin_risks(state.project)
    except Exception as e:
        state.audit["risk_registry_parse_error"] = f"{e}: {raw[:300]}"
        state.audit["risk_registry_used_builtin"] = True
        data = _builtin_risks(state.project)

    return _fill_defaults(data, state.project)


# =============================================================================
# SECTION 3 — MARKDOWN (guardrails-valid)
# =============================================================================

def _build_markdown(data: Dict[str, Any], proj, standards: dict) -> str:
    risks    = data.get("risks", [])
    sections = standards["docs"]["risk_registry"].get("required_sections", [])
    org      = standards["org"]["name"]
    lines: List[str] = []

    def h(t):  lines.extend([f"## {t}", ""])
    def b(t):  lines.extend([t, ""])
    def li(t): lines.append(f"- {t}")

    for sec in sections:
        sl = sec.lower()
        h(sec)

        if sl == "overview":
            b(f"This Risk Registry and Risk Assessment is maintained for {proj.project_name} "
              f"(Ref: {proj.project_id}) in accordance with {org} PMO governance standards.")
            b(data.get("risk_framework", ""))

        elif sl == "risk summary":
            b(data.get("executive_summary", ""))
            counts: Dict[str, int] = {}
            for r in risks:
                counts[r.get("rating","Medium")] = counts.get(r.get("rating","Medium"),0)+1
            li(f"Total risks identified: {len(risks)}")
            for level in ("Critical","High","Medium","Low"):
                if counts.get(level):
                    li(f"{level}: {counts[level]} risk(s)")
            lines.append("")
            li(f"Overall risk posture: {data.get('risk_posture','ELEVATED')}")
            lines.append("")

        elif sl == "detailed risks":
            for r in risks:
                li(f"{r['id']} [{r.get('source','AI Identified')}] {r['title']}: {r['description'][:120]}")
            lines.append("")

        elif sl == "mitigations":
            for r in risks:
                li(f"{r['id']} ({r['title']}): {r['mitigation'][:150]}")
            lines.append("")

        elif sl == "owners":
            li("All risks are assigned to the Project Manager as default owner.")
            li("The PM reviews and reassigns each risk to the appropriate domain specialist.")
            for r in risks:
                note = r.get("owner_note","")
                li(f"{r['id']} — Owner: Project Manager | {note}")
            lines.append("")

        elif sl == "registry overview":
            cats = len(set(r["category"] for r in risks))
            b(f"This registry documents {len(risks)} risks across {cats} categories. "
              f"All risks scored on a 5x5 likelihood-impact matrix per ISO 31000.")
            b("The register is version-controlled and retained for 7 years minimum.")

        elif sl == "risk list":
            for r in risks:
                li(f"{r['id']} | {r['category']} | {r['title']} | "
                   f"L:{r.get('likelihood_score',3)} x I:{r.get('impact_score',3)} = {r['score']} "
                   f"({r['rating']}) | Owner: Project Manager | Status: {r['status']}")
            lines.append("")

        elif sl == "review cadence":
            for c in data.get("review_cadence", []):
                li(f"{c.get('review_type','')}: {c.get('frequency','')} | "
                   f"Attendees: {c.get('attendees','')} | Output: {c.get('output','')}")
            lines.append("")

        elif sl == "approvals":
            b("Approval signatures required before this document is ratified.")
            li("Project Sponsor | Name: _________________ | Signature: ________ | Date: ________")
            li("Project Manager | Name: _________________ | Signature: ________ | Date: ________")
            li("PMO Lead | Name: _________________ | Signature: ________ | Date: ________")
            li("Chief Risk Officer | Name: _________________ | Signature: ________ | Date: ________")
            lines.append("")

        else:
            b(f"See detailed section for '{sec}' content.")
            for r in risks[:5]:
                li(f"{r['id']}: {r['title']}")
            lines.append("")

    h("Assumptions")
    for a in data.get("assumptions", []):
        li(a)
    lines.append("")

    return "\n".join(lines)


# =============================================================================
# SECTION 4 — WORD DOCUMENT RENDERER
# =============================================================================

# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _shd(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"),   "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"),  hex_color)
    tcPr.append(s)


def _borders(table, color: str = "C5D3E8") -> None:
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcB  = OxmlElement("w:tcBorders")
            for side in ("top","left","bottom","right","insideH","insideV"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"),   "single")
                b.set(qn("w:sz"),    "4")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), color)
                tcB.append(b)
            # tcBorders must come before w:shd in tcPr (OOXML schema order).
            # Insert at position 1 (after tcW which is at 0).
            shd_els = tcPr.findall(qn("w:shd"))
            if shd_els:
                # Insert tcBorders just before the first shd element
                shd_idx = list(tcPr).index(shd_els[0])
                tcPr.insert(shd_idx, tcB)
            else:
                shd_els = tcPr.findall(qn("w:shd"))
            if shd_els:
                tcPr.insert(list(tcPr).index(shd_els[0]), tcB)
            else:
                tcPr.append(tcB)


def _cw(cell, cm: float) -> None:
    """Set cell width — replaces any existing tcW to avoid duplicate elements."""
    tcPr = cell._tc.get_or_add_tcPr()
    # Remove any existing tcW elements first
    for existing in tcPr.findall(qn("w:tcW")):
        tcPr.remove(existing)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    # tcW must be the first child of tcPr per OOXML schema
    tcPr.insert(0, tcW)


def _page_num(run) -> None:
    for tag, txt in (("begin", None), ("instrText", "PAGE"), ("end", None)):
        if tag == "instrText":
            el = OxmlElement("w:instrText"); el.text = txt
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
        run._r.append(el)


# ── Risk colour helpers ───────────────────────────────────────────────────────

def _rgb(level: str) -> RGBColor:
    l = level.upper()
    if l == "CRITICAL": return _RED
    if l == "HIGH":     return RGBColor(0xB4, 0x53, 0x09)
    if l == "MEDIUM":   return _AMBER
    return _GREEN


def _bg(level: str) -> str:
    l = level.upper()
    if l == "CRITICAL": return _CRIT_BG
    if l == "HIGH":     return _RED_BG
    if l == "MEDIUM":   return _AMBER_BG
    return _GREEN_BG


# ── Document setup ────────────────────────────────────────────────────────────

def _setup(doc: Document) -> None:
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9.5)
    doc.styles["Normal"].font.color.rgb = _BODY
    s = doc.sections[0]
    s.page_width    = Cm(21)
    s.page_height   = Cm(29.7)
    s.left_margin   = Cm(1.6)
    s.right_margin  = Cm(1.6)
    s.top_margin    = Cm(1.8)
    s.bottom_margin = Cm(1.6)
    # Fix w:zoom — add required w:percent attribute to satisfy OOXML schema validation
    for zoom in doc.settings.element.findall(qn("w:zoom")):
        if zoom.get(qn("w:percent")) is None:
            zoom.set(qn("w:percent"), "100")


def _footer(doc: Document, project_id: str, org: str, date: str) -> None:
    ft = doc.sections[0].footer
    ft.is_linked_to_previous = False
    p  = ft.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(
        f"{org}  ·  Risk Registry & Risk Assessment  ·  {project_id}  ·  {date}  ·  Page "
    )
    r1.font.size = Pt(7.5); r1.font.color.rgb = _MUTED
    r2 = p.add_run()
    r2.font.size = Pt(7.5); r2.font.color.rgb = _MUTED
    _page_num(r2)


# ── Typography primitives ─────────────────────────────────────────────────────

def _sec_head(doc: Document, text: str) -> None:
    """Navy bold section heading with blue underline. Tight spacing — no gaps."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text.upper())
    r.bold = True; r.font.size = Pt(10.5)
    r.font.color.rgb = _NAVY; r.font.name = "Calibri"
    # pBdr must appear early in pPr (after pStyle/numPr, before shd/tabs).
    # Insert at index 0 so it precedes spacing elements added by python-docx.
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "0072BB")
    pBdr.append(bot)
    pPr.insert(0, pBdr)


def _sub(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(9.5)
    r.font.color.rgb = _BLUE; r.font.name = "Calibri"


def _body(doc: Document, text: str, italic: bool = False,
          color: RGBColor = None, size: float = 9.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = color or _BODY
    r.italic = italic; r.font.name = "Calibri"


# ── Table cell primitives ─────────────────────────────────────────────────────

def _hc(cell, text: str, w: float = 0, bg: str = _HDR_BG, sz: float = 8.5) -> None:
    """Header cell — dark bg, white bold text."""
    if w: _cw(cell, w)
    _shd(cell, bg)
    p = cell.paragraphs[0]; p.clear()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(sz)
    r.font.color.rgb = _WHITE; r.font.name = "Calibri"


def _dc(cell, text: str, w: float = 0, bg: str = _ROW_A,
        color: RGBColor = None, bold: bool = False, sz: float = 8.5,
        align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Data cell."""
    if w: _cw(cell, w)
    _shd(cell, bg)
    p = cell.paragraphs[0]; p.clear()
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(str(text or "—"))
    r.font.size = Pt(sz); r.font.color.rgb = color or _BODY
    r.bold = bold; r.font.name = "Calibri"


# =============================================================================
# SECTION RENDERERS  — no page breaks between them; sections flow continuously
# =============================================================================

def _cover(doc: Document, proj, standards: dict,
           date: str, logo_path: Optional[str]) -> None:
    org    = standards["org"]["name"]
    header = standards["org"].get("doc_header", "Internal")

    # Banner bar
    bar = doc.add_table(rows=1, cols=2)
    bar.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc = bar.cell(0,0); _shd(lc, _HDR_BG); _cw(lc, 5.0)
    pl = lc.paragraphs[0]
    pl.paragraph_format.space_before = Pt(8)
    pl.paragraph_format.space_after  = Pt(8)
    if logo_path and os.path.exists(logo_path):
        pl.add_run().add_picture(logo_path, width=Cm(4.0))
    else:
        r = pl.add_run(org.upper())
        r.bold = True; r.font.size = Pt(16)
        r.font.color.rgb = _WHITE; r.font.name = "Calibri"

    rc = bar.cell(0,1); _shd(rc, _HDR_BG); _cw(rc, 12.4)
    pr = rc.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr.paragraph_format.space_before = Pt(12)
    pr.paragraph_format.space_after  = Pt(12)
    r2 = pr.add_run(f"GOVERNANCE & RISK MANAGEMENT  |  {header.upper()}")
    r2.bold = True; r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor(0xCC,0xE4,0xFF); r2.font.name = "Calibri"

    # Title block
    def _tp(txt, size, color, before=0, after=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)
        r = p.add_run(txt)
        r.bold = True; r.font.size = Pt(size)
        r.font.color.rgb = color; r.font.name = "Calibri"

    _tp("PROGRAMME RISK DOCUMENTATION", 8.5, _MUTED, before=18, after=4)
    _tp("Risk Registry", 28, _NAVY, before=0, after=0)
    _tp("& Risk Assessment", 28, _BLUE, before=0, after=6)
    _tp(proj.project_name, 13, _MID, before=0, after=10)

    # Metadata table
    meta = [
        ("Project Reference",  proj.project_id),
        ("Project Type",       proj.project_type),
        ("Sponsor",            proj.sponsor or "To be confirmed"),
        ("Organisation",       org),
        ("Classification",     header),
        ("Document Version",   "v1.0 — Initial Release"),
        ("Prepared By",        f"{org} PMO Risk Team"),
        ("Document Date",      date),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.style = "Table Grid"
    for i,(k,v) in enumerate(meta):
        _hc(tbl.rows[i].cells[0], k, 4.5, sz=9.0)
        _dc(tbl.rows[i].cells[1], v, 12.9, _ROW_A, sz=9.0)
    _borders(tbl)

    p_conf = doc.add_paragraph()
    p_conf.paragraph_format.space_before = Pt(10)
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_conf = p_conf.add_run(
        "COMMERCIAL-IN-CONFIDENCE  |  FOR INTERNAL USE ONLY  |  DO NOT DISTRIBUTE"
    )
    r_conf.font.size = Pt(7.5); r_conf.italic = True
    r_conf.font.color.rgb = _MUTED; r_conf.font.name = "Calibri"

    doc.add_page_break()  # Only page break in the entire document — after cover


def _exec_summary(doc: Document, data: dict, risks: list) -> None:
    _sec_head(doc, "1.  Executive Summary")
    _body(doc, data.get("executive_summary", ""))

    counts = {"Critical":0,"High":0,"Medium":0,"Low":0}
    for r in risks:
        lvl = r.get("rating","Low")
        if lvl in counts: counts[lvl] += 1

    ai_c    = sum(1 for r in risks if r.get("source")=="AI Identified")
    human_c = sum(1 for r in risks if r.get("source")=="Human Identified")

    tbl = doc.add_table(rows=2, cols=6)
    tbl.style = "Table Grid"
    hdrs  = ["Total","Critical","High","Medium","Low","Risk Posture"]
    vals  = [str(len(risks)), str(counts["Critical"]), str(counts["High"]),
             str(counts["Medium"]), str(counts["Low"]),
             data.get("risk_posture","ELEVATED")]
    vcols = [_BODY, _RED, _rgb("HIGH"), _AMBER, _GREEN,
             _RED if data.get("risk_posture")=="ELEVATED" else _AMBER]
    vbgs  = [_ROW_A, _CRIT_BG, _RED_BG, _AMBER_BG, _GREEN_BG,
             _CRIT_BG if data.get("risk_posture")=="ELEVATED" else _AMBER_BG]
    w = 2.9
    for i,(h,v,vc,vb) in enumerate(zip(hdrs,vals,vcols,vbgs)):
        _hc(tbl.rows[0].cells[i], h, w)
        _dc(tbl.rows[1].cells[i], v, w, vb, vc, True, 14.0,
            WD_ALIGN_PARAGRAPH.CENTER)
    _borders(tbl)

    if human_c > 0 and ai_c > 0:
        src_note = (f"This register contains {human_c} risk(s) formally raised by the project team "
                    f"and {ai_c} additional risk(s) identified through structured PMO analysis "
                    f"of the project scope, deliverables, and operating environment. "
                    f"All risks are owned by the Project Manager pending formal reassignment.")
    elif human_c > 0:
        src_note = (f"All {human_c} risks in this register were formally raised by the project team. "
                    f"All risks are owned by the Project Manager pending formal reassignment.")
    else:
        src_note = (f"All {ai_c} risks were identified through structured PMO analysis of the project "
                    f"scope, deliverables, and operating environment. "
                    f"All risks are owned by the Project Manager pending formal reassignment.")
    _body(doc, src_note, italic=True, color=_MUTED, size=8.5)


def _framework(doc: Document, data: dict) -> None:
    _sec_head(doc, "2.  Risk Framework & Methodology")
    _body(doc, data.get("risk_framework",""))
    _body(doc, ("All risks are initially owned by the Project Manager. The PM reviews each risk, "
                "confirms or adjusts the rating, and reassigns to the appropriate domain specialist "
                "as indicated in the Ownership Register (Section 7). "
                "Human-identified risks are those raised explicitly by the project team and are "
                "flagged in the register. AI-identified risks are generated by automated analysis "
                "of the project description and should be validated by the project team."))

    _sub(doc, "Risk Scoring Matrix  (Likelihood x Impact, 5x5 scale)")
    likelihoods = [("Almost Certain (5)",5),("Likely (4)",4),("Possible (3)",3),
                   ("Unlikely (2)",2),("Rare (1)",1)]
    impacts     = [("Negligible (1)",1),("Minor (2)",2),("Moderate (3)",3),
                   ("Major (4)",4),("Catastrophic (5)",5)]
    tbl = doc.add_table(rows=6, cols=6)
    tbl.style = "Table Grid"
    _hc(tbl.rows[0].cells[0], "Likelihood \\ Impact", 3.6)
    for j,(lbl,_) in enumerate(impacts):
        _hc(tbl.rows[0].cells[j+1], lbl, 2.2)
    for i,(llbl,lv) in enumerate(likelihoods):
        _hc(tbl.rows[i+1].cells[0], llbl, 3.6)
        for j,(_,iv) in enumerate(impacts):
            score = lv*iv
            sbg  = _CRIT_BG if score>=16 else (_RED_BG if score>=9 else
                   (_AMBER_BG if score>=4 else _GREEN_BG))
            scol = _RED if score>=16 else (_rgb("HIGH") if score>=9 else
                   (_AMBER if score>=4 else _GREEN))
            _dc(tbl.rows[i+1].cells[j+1], str(score), 2.2, sbg, scol,
                score>=9, 8.5, WD_ALIGN_PARAGRAPH.CENTER)
    _borders(tbl)

    _sub(doc, "Rating Thresholds")
    leg = [("16–25","CRITICAL",_CRIT_BG,_RED,"Immediate — Sponsor + Steering Committee","24 hours"),
           ("9–15","HIGH",_RED_BG,_rgb("HIGH"),"PM + Sponsor within 48 hours","Weekly"),
           ("4–8","MEDIUM",_AMBER_BG,_AMBER,"Programme Manager review","Fortnightly"),
           ("1–3","LOW",_GREEN_BG,_GREEN,"Risk Owner monitoring","Monthly")]
    ltbl = doc.add_table(rows=5, cols=4)
    ltbl.style = "Table Grid"
    for i,h in enumerate(("Score","Rating","Escalation Requirement","Review Frequency")):
        _hc(ltbl.rows[0].cells[i], h, [1.8,2.0,9.6,3.0][i])
    for i,(score,rating,rbg,rcol,esc,freq) in enumerate(leg):
        _dc(ltbl.rows[i+1].cells[0], score,  1.8, rbg)
        _dc(ltbl.rows[i+1].cells[1], rating, 2.0, rbg, rcol, True)
        _dc(ltbl.rows[i+1].cells[2], esc,    9.6, rbg)
        _dc(ltbl.rows[i+1].cells[3], freq,   3.0, rbg)
    _borders(ltbl)


def _assumptions(doc: Document, data: dict) -> None:
    _sec_head(doc, "3.  Risk Assumptions Register")
    _body(doc, ("The assumptions below underpin all risk assessments in this register. "
                "If any assumption is invalidated during delivery, affected risks must be "
                "reassessed immediately and escalated to the Project Manager within 48 hours."))

    assumptions = data.get("assumptions", [])
    for i, a in enumerate(assumptions):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.4)
        ref_run = p.add_run(f"A-{str(i+1).zfill(3)}  ")
        ref_run.bold = True
        ref_run.font.color.rgb = _BLUE
        ref_run.font.size = Pt(9.0)
        ref_run.font.name = "Calibri"
        txt_run = p.add_run(a)
        txt_run.font.size = Pt(9.0)
        txt_run.font.color.rgb = _BODY
        txt_run.font.name = "Calibri"


def _summary_dashboard(doc: Document, risks: list) -> None:
    _sec_head(doc, "4.  Risk Summary Dashboard")
    _body(doc, (f"The table below consolidates all {len(risks)} identified risks for this project. "
                "Each risk is currently owned by the Project Manager. The 'Reassign To' column "
                "identifies the recommended specialist for each risk domain. Risks are ordered by "
                "composite score — highest exposure first."))
    risks = sorted(risks, key=lambda r: r.get("score", 0), reverse=True)

    cols   = ("ID","Source","Category","Risk Title","L","I","Score","Rating","Reassign To","Status")
    widths = (1.1, 2.0, 2.0, 4.2, 0.6, 0.6, 0.8, 1.8, 3.8, 1.7)
    tbl = doc.add_table(rows=len(risks)+1, cols=len(cols))
    tbl.style = "Table Grid"
    for i,(h,w) in enumerate(zip(cols,widths)):
        _hc(tbl.rows[0].cells[i], h, w)

    for i,r in enumerate(risks):
        bg = _ROW_A if i%2==0 else _ROW_B
        src_bg = "E8F4FD" if r.get("source")=="AI Identified" else "FFF8E1"
        note = r.get("owner_note","—")
        if note.lower().startswith("reassign to:"):
            note = note[len("reassign to:"):].strip()
        vals = [
            (r["id"],       bg,             _BLUE,          True,  7.5),
            (r.get("source","AI"), src_bg,  _MID,           False, 7.0),
            (r["category"], bg,             _BODY,          False, 7.5),
            (r["title"],    bg,             _BODY,          False, 7.5),
            (str(r.get("likelihood_score",3)), _bg(r["rating"]), _rgb(r["rating"]), True, 8.5),
            (str(r.get("impact_score",3)),     _bg(r["rating"]), _rgb(r["rating"]), True, 8.5),
            (str(r["score"]),                  _bg(r["rating"]), _rgb(r["rating"]), True, 8.5),
            (r["rating"],                      _bg(r["rating"]), _rgb(r["rating"]), True, 8.0),
            (note,          bg,             _MID,           False, 7.0),
            (r["status"],   bg,             _BODY,          False, 7.5),
        ]
        for j,((val,vbg,vcol,vbold,vsz),w) in enumerate(zip(vals,widths)):
            al = WD_ALIGN_PARAGRAPH.CENTER if j in (4,5,6,7) else WD_ALIGN_PARAGRAPH.LEFT
            _dc(tbl.rows[i+1].cells[j], val, w, vbg, vcol, vbold, vsz, al)
    _borders(tbl)


def _detailed_risks(doc: Document, risks: list) -> None:
    _sec_head(doc, "5.  Risk Registry \u2014 Detailed Entries")
    _body(doc, (f"The following {len(risks)} risk entries constitute the full Risk Registry for this project. "
                "Each entry shows the risk narrative, scoring metadata, mitigation actions, and contingency plan. "
                "The Project Manager holds ownership of all risks and must formally reassign each risk to the "
                "appropriate domain specialist within 5 business days of document approval."))
    risks = sorted(risks, key=lambda r: r.get("score", 0), reverse=True)

    for idx, r in enumerate(risks):
        # \u2500\u2500 Coloured header bar per risk \u2500\u2500
        hdr_bg = {"CRITICAL": "8B0000", "HIGH": "7B2D00",
                  "MEDIUM": "7A5800", "LOW": "1B5E20"}.get(r["rating"].upper(), _HDR_BG)
        hdr = doc.add_table(rows=1, cols=4)
        hdr.style = "Table Grid"
        _hc(hdr.rows[0].cells[0], r["id"],                            1.4, hdr_bg, 10.0)
        _hc(hdr.rows[0].cells[1], f"{r['title']}  [{r['category']}]", 9.9, hdr_bg, 10.0)
        _hc(hdr.rows[0].cells[2], f"Score: {r['score']}  ({r['rating']})", 2.8, hdr_bg, 10.0)
        src_badge = "1A3A5C" if r.get("source") == "AI Identified" else "4A3000"
        _hc(hdr.rows[0].cells[3], r.get("source", "AI Identified"),   3.5, src_badge, 8.5)
        _borders(hdr, hdr_bg)

        # \u2500\u2500 Description (prose) \u2500\u2500
        desc_p = doc.add_paragraph()
        desc_p.paragraph_format.space_before = Pt(4)
        desc_p.paragraph_format.space_after  = Pt(2)
        desc_p.paragraph_format.left_indent  = Cm(0.3)
        lbl1 = desc_p.add_run("Description:  ")
        lbl1.bold = True; lbl1.font.size = Pt(9.0); lbl1.font.color.rgb = _NAVY; lbl1.font.name = "Calibri"
        txt1 = desc_p.add_run(r["description"])
        txt1.font.size = Pt(9.0); txt1.font.color.rgb = _BODY; txt1.font.name = "Calibri"

        # \u2500\u2500 Trigger (prose) \u2500\u2500
        trig_p = doc.add_paragraph()
        trig_p.paragraph_format.space_before = Pt(2)
        trig_p.paragraph_format.space_after  = Pt(4)
        trig_p.paragraph_format.left_indent  = Cm(0.3)
        lbl2 = trig_p.add_run("Risk Trigger:  ")
        lbl2.bold = True; lbl2.font.size = Pt(9.0); lbl2.font.color.rgb = _NAVY; lbl2.font.name = "Calibri"
        txt2 = trig_p.add_run(r.get("trigger", "\u2014"))
        txt2.font.size = Pt(9.0); txt2.font.color.rgb = _MID; txt2.font.name = "Calibri"

        # \u2500\u2500 Compact metadata table (L / I / Score / Rating / Review / Owner / Reassign / Residual) \u2500\u2500
        note = r.get("owner_note", "\u2014")
        if note.lower().startswith("reassign to:"):
            note = note[len("reassign to:"):].strip()
        meta_cols   = ("Likelihood", "Impact", "Score", "Rating", "Review", "Owner", "Reassign To", "Residual")
        meta_widths = (2.2, 2.0, 1.4, 1.8, 1.9, 3.0, 3.5, 1.8)
        meta_vals   = (
            f"{r['likelihood']} ({r.get('likelihood_score', '\u2014')})",
            f"{r['impact']} ({r.get('impact_score', '\u2014')})",
            str(r["score"]),
            r["rating"],
            r.get("review_frequency", "Monthly"),
            "Project Manager",
            note,
            r.get("residual_rating", "\u2014"),
        )
        mtbl = doc.add_table(rows=2, cols=len(meta_cols))
        mtbl.style = "Table Grid"
        for ci, (h, w) in enumerate(zip(meta_cols, meta_widths)):
            _hc(mtbl.rows[0].cells[ci], h, w, sz=8.0)
        for ci, (v, w) in enumerate(zip(meta_vals, meta_widths)):
            col_name = meta_cols[ci]
            if col_name == "Rating":
                _dc(mtbl.rows[1].cells[ci], v, w, _bg(r["rating"]), _rgb(r["rating"]), True, 8.5, WD_ALIGN_PARAGRAPH.CENTER)
            elif col_name == "Score":
                _dc(mtbl.rows[1].cells[ci], v, w, _bg(r["rating"]), _rgb(r["rating"]), True, 9.0, WD_ALIGN_PARAGRAPH.CENTER)
            elif col_name == "Residual":
                _dc(mtbl.rows[1].cells[ci], v, w, _bg(r.get("residual_rating", "Low")), _rgb(r.get("residual_rating", "Low")), True, 8.5, WD_ALIGN_PARAGRAPH.CENTER)
            elif col_name in ("Owner", "Reassign To"):
                _dc(mtbl.rows[1].cells[ci], v, w, "FFF8E1", _AMBER, False, 8.0)
            else:
                _dc(mtbl.rows[1].cells[ci], v, w, _ROW_A, _BODY, False, 8.0)
        _borders(mtbl)

        # \u2500\u2500 Mitigation Actions (label + bullets) \u2500\u2500
        mit_p = doc.add_paragraph()
        mit_p.paragraph_format.space_before = Pt(6)
        mit_p.paragraph_format.space_after  = Pt(1)
        mit_p.paragraph_format.left_indent  = Cm(0.3)
        m_lbl = mit_p.add_run("Mitigation Actions")
        m_lbl.bold = True; m_lbl.font.size = Pt(9.0); m_lbl.font.color.rgb = _NAVY; m_lbl.font.name = "Calibri"

        for step in r["mitigation"].split("."):
            step = step.strip().strip(";")
            if len(step) < 6:
                continue
            bp = doc.add_paragraph()
            bp.paragraph_format.space_before = Pt(1)
            bp.paragraph_format.space_after  = Pt(1)
            bp.paragraph_format.left_indent  = Cm(0.8)
            bullet = bp.add_run(f"\u2022  {step}.")
            bullet.font.size = Pt(8.5); bullet.font.color.rgb = _BODY; bullet.font.name = "Calibri"

        # \u2500\u2500 Contingency Plan (label + bullets) \u2500\u2500
        con_p = doc.add_paragraph()
        con_p.paragraph_format.space_before = Pt(5)
        con_p.paragraph_format.space_after  = Pt(1)
        con_p.paragraph_format.left_indent  = Cm(0.3)
        c_lbl = con_p.add_run("Contingency Plan")
        c_lbl.bold = True; c_lbl.font.size = Pt(9.0); c_lbl.font.color.rgb = _rgb(r["rating"]); c_lbl.font.name = "Calibri"

        for step in r["contingency"].split("."):
            step = step.strip().strip(";")
            if len(step) < 6:
                continue
            bp = doc.add_paragraph()
            bp.paragraph_format.space_before = Pt(1)
            bp.paragraph_format.space_after  = Pt(1)
            bp.paragraph_format.left_indent  = Cm(0.8)
            bullet = bp.add_run(f"\u25ba  {step}.")
            bullet.font.size = Pt(8.5); bullet.font.color.rgb = _MID; bullet.font.name = "Calibri"

        # \u2500\u2500 Spacer between risk cards \u2500\u2500
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(8)
        spacer.paragraph_format.space_after  = Pt(0)


def _mitigation_matrix(doc: Document, risks: list) -> None:
    _sec_head(doc, "6.  Risk vs. Mitigation Matrix")
    _body(doc, (f"The matrix below maps each of the {len(risks)} identified project risks "
                "directly to its mitigation strategy, contingency plan, and post-mitigation "
                "residual rating. This is the primary treatment reference for the Project Manager, "
                "Steering Committee, and internal audit. Mitigation actions must be actioned "
                "by the designated risk owner within the review frequency stated for each risk. "
                "Contingency plans are pre-approved and may be activated by the Project Manager "
                "for Medium risks, and by the Programme Sponsor for High and Critical risks."))
    risks = sorted(risks, key=lambda r: r.get("score", 0), reverse=True)

    cols   = ("ID","Risk Title","Rating","Mitigation Strategy","Contingency Plan","Residual")
    widths = (1.1, 3.0, 1.8, 5.8, 4.5, 1.4)
    tbl = doc.add_table(rows=len(risks)+1, cols=len(cols))
    tbl.style = "Table Grid"
    for i,(h,w) in enumerate(zip(cols,widths)):
        _hc(tbl.rows[0].cells[i], h, w)
    for i,r in enumerate(risks):
        bg = _ROW_A if i%2==0 else _ROW_B
        _dc(tbl.rows[i+1].cells[0], r["id"],        1.1, bg, _BLUE, True, 8.0)
        _dc(tbl.rows[i+1].cells[1], r["title"],      3.0, bg, _BODY, False, 7.5)
        _dc(tbl.rows[i+1].cells[2], r["rating"],     1.8, _bg(r["rating"]),
            _rgb(r["rating"]), True, 8.0, WD_ALIGN_PARAGRAPH.CENTER)
        _dc(tbl.rows[i+1].cells[3], r["mitigation"], 5.8, bg, _BODY, False, 7.5)
        _dc(tbl.rows[i+1].cells[4], r["contingency"],4.5, bg, _BODY, False, 7.5)
        _dc(tbl.rows[i+1].cells[5], r["residual_rating"], 1.4,
            _bg(r["residual_rating"]), _rgb(r["residual_rating"]),
            True, 8.0, WD_ALIGN_PARAGRAPH.CENTER)
    _borders(tbl)


def _ownership_register(doc: Document, risks: list) -> None:
    _sec_head(doc, "7.  Risk Ownership Register")
    _body(doc, (f"The Project Manager holds initial ownership of all {len(risks)} risks in this "
                "register. This is a governance requirement to ensure no risk is unowned at "
                "point of identification. The PM must review each risk, validate the rating "
                "and proposed mitigation, and formally reassign ownership to the appropriate "
                "domain specialist. Reassignment must be recorded in the RAID log and "
                "communicated to the incoming owner in writing."))

    cols   = ("Risk ID","Risk Title","Risk Category","Current Owner","Recommended Owner","Rating","Review Freq.")
    widths = (1.2, 4.0, 2.2, 3.0, 3.6, 1.6, 2.0)
    tbl = doc.add_table(rows=len(risks)+1, cols=len(cols))
    tbl.style = "Table Grid"
    for i,(h,w) in enumerate(zip(cols,widths)):
        _hc(tbl.rows[0].cells[i], h, w)
    for i,r in enumerate(risks):
        bg = _ROW_A if i%2==0 else _ROW_B
        note = r.get("owner_note","—")
        if note.lower().startswith("reassign to:"):
            note = note[len("reassign to:"):].strip()
        _dc(tbl.rows[i+1].cells[0], r["id"],              1.2, bg, _BLUE, True)
        _dc(tbl.rows[i+1].cells[1], r["title"],            4.0, bg)
        _dc(tbl.rows[i+1].cells[2], r.get("category","—"), 2.2, bg, _MID)
        _dc(tbl.rows[i+1].cells[3], "Project Manager",     3.0, "FFF8E1", _AMBER, True)
        _dc(tbl.rows[i+1].cells[4], note,                  3.6, bg, _NAVY)
        _dc(tbl.rows[i+1].cells[5], r["rating"],           1.6,
            _bg(r["rating"]), _rgb(r["rating"]), True, 8.0, WD_ALIGN_PARAGRAPH.CENTER)
        _dc(tbl.rows[i+1].cells[6], r.get("review_frequency","Monthly"), 2.0, bg)
    _borders(tbl)
    _body(doc, ("Note: The Project Manager retains full accountability for all risks until "
                "formal reassignment is documented and acknowledged by the incoming risk owner "
                "in writing. Any risk rated High or Critical must be reassigned within "
                "48 hours of this document being approved."),
          italic=True, color=_MUTED, size=8.0)


def _review_cadence(doc: Document, data: dict) -> None:
    _sec_head(doc, "8.  Review Cadence & Governance")
    _body(doc, ("Risk reviews conducted at the cadence below. All outcomes logged in the programme "
                "audit trail and retained for 7 years per document retention policy."))

    cadence = data.get("review_cadence", [])
    cols    = ("Review Type","Frequency","Attendees","Scope","Output Document")
    widths  = (3.6, 2.2, 4.2, 3.8, 3.8)
    tbl = doc.add_table(rows=len(cadence)+1, cols=len(cols))
    tbl.style = "Table Grid"
    for i,(h,w) in enumerate(zip(cols,widths)):
        _hc(tbl.rows[0].cells[i], h, w)
    for i,c in enumerate(cadence):
        bg = _ROW_A if i%2==0 else _ROW_B
        _dc(tbl.rows[i+1].cells[0], c.get("review_type","—"), 3.6, bg, _NAVY, True)
        _dc(tbl.rows[i+1].cells[1], c.get("frequency","—"),   2.2, bg)
        _dc(tbl.rows[i+1].cells[2], c.get("attendees","—"),   4.2, bg)
        _dc(tbl.rows[i+1].cells[3], c.get("scope","—"),       3.8, bg)
        _dc(tbl.rows[i+1].cells[4], c.get("output","—"),      3.8, bg, _BLUE)
    _borders(tbl)
    _body(doc, ("Critical risks must be reviewed within 24 hours of a rating change. "
                "PM notifies Sponsor immediately upon any risk reaching Critical status."),
          italic=True, color=_MUTED, size=8.0)


def _approvals(doc: Document, standards: dict) -> None:
    _sec_head(doc, "9.  Approvals & Sign-off")
    _body(doc, ("By signing below, each approver confirms they have reviewed this Risk Registry, "
                "accept the risk ratings and default PM ownership, and authorise the documented "
                "mitigation and contingency strategies."))

    approvers = [
        ("Project Sponsor",               "Overall programme accountability"),
        ("Project Manager",               "Risk register ownership and reassignment authority"),
        ("PMO Representative",            "Governance and compliance verification"),
        ("Chief Risk Officer",            "Risk methodology and escalation threshold approval"),
        ("Head of Information Security",  "Security and data risk sign-off"),
        ("Internal Audit Representative", "Independent assurance and completeness check"),
    ]
    cols   = ("Role","Accountability","Name","Signature","Date","Ver.")
    widths = (3.8, 4.4, 3.3, 3.2, 1.8, 1.1)
    tbl = doc.add_table(rows=len(approvers)+1, cols=len(cols))
    tbl.style = "Table Grid"
    for i,(h,w) in enumerate(zip(cols,widths)):
        _hc(tbl.rows[0].cells[i], h, w)
    for i,(role,acc) in enumerate(approvers):
        bg = _ROW_A if i%2==0 else _ROW_B
        _dc(tbl.rows[i+1].cells[0], role, 3.8, bg, _BODY, True)
        _dc(tbl.rows[i+1].cells[1], acc,  4.4, bg, _MID)
        _dc(tbl.rows[i+1].cells[2], "",   3.3, bg)
        _dc(tbl.rows[i+1].cells[3], "",   3.2, bg)
        _dc(tbl.rows[i+1].cells[4], "",   1.8, bg)
        _dc(tbl.rows[i+1].cells[5], "v1.0", 1.1, bg, _MUTED, False, 7.5,
            WD_ALIGN_PARAGRAPH.CENTER)
    _borders(tbl)

    org = standards["org"]["name"]
    _body(doc, (f"Classification: Commercial-in-Confidence  ·  Retention: 7 years  ·  "
                f"Owner: {org} PMO Risk Team  ·  Re-approval required for any amendment."),
          italic=True, color=_MUTED, size=7.5)


# =============================================================================
# SECTION 5 — MAIN DOCX ASSEMBLER
# =============================================================================

def _resolve_logo(standards: dict) -> Optional[str]:
    org_name = standards["org"].get("name","").lower().replace(" ","")
    candidates = [
        "nttdata_logo.png", "logo.png",
        os.path.join("config","logo.png"),
        os.path.join("assets","logo.png"),
        f"{org_name}_logo.png",
    ]
    for c in candidates:
        full = os.path.join(os.path.dirname(os.path.abspath(__file__)), c)
        if os.path.exists(full): return full
        if os.path.exists(c):    return c
    return None


def build_risk_registry_docx(data: Dict[str, Any], state: PMOState,
                              logo_path: Optional[str] = None) -> bytes:
    """
    Assemble the complete Word document from structured risk data.
    Returns raw bytes for HTTP response or file write.

    Layout rules:
      - ONE page break total: after the cover page.
      - All body sections flow continuously. _sec_head() provides 12pt
        space_before which creates visual separation without blank pages.
      - Tables use tight cell padding (2pt) for dense, professional layout.
    """
    doc   = Document()
    _setup(doc)

    proj  = state.project
    org   = state.standards["org"]["name"]
    date  = datetime.now().strftime("%d %B %Y")
    risks = data.get("risks", [])

    if logo_path is None:
        logo_path = _resolve_logo(state.standards)

    _footer(doc, proj.project_id, org, date)

    # Cover (ends with page break — only page break in the document)
    _cover(doc, proj, state.standards, date, logo_path)

    # Body sections — continuous flow, no page breaks
    _exec_summary(doc, data, risks)
    _framework(doc, data)
    _assumptions(doc, data)
    _summary_dashboard(doc, risks)
    _detailed_risks(doc, risks)
    _mitigation_matrix(doc, risks)
    _ownership_register(doc, risks)
    _review_cadence(doc, data)
    _approvals(doc, state.standards)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =============================================================================
# SECTION 6 — PUBLIC API
# =============================================================================

def generate_risk_registry_artifact(state: PMOState) -> DocumentArtifact:
    """
    Full pipeline:
      1. LLM analyses project description → structured JSON with auto-identified risks
      2. Guardrail-valid markdown built from JSON
      3. Professional Word document rendered
      4. DocumentArtifact returned; docx bytes cached in state.audit["risk_registry_docx"]
    """
    doc_title = state.standards["docs"]["risk_registry"].get("title","Risk Registry")

    # Ensure raw_upload_text (from project documents uploaded by the user)
    # is always visible to the risk prompt — even if not explicitly mapped
    # to "risk_registry" in the UI. _build_prompt reads both sources.
    state.audit.setdefault("uploaded_mapping", {})

    try:
        risk_data = _call_llm(state)
        state.audit["risk_registry_llm_ok"] = True
        state.audit["risk_registry_structured_data"] = risk_data
    except Exception as e:
        state.audit["risk_registry_llm_error"] = str(e)
        risk_data = _fill_defaults(_builtin_risks(state.project), state.project)

    md = _build_markdown(risk_data, state.project, state.standards)

    try:
        docx_bytes = build_risk_registry_docx(risk_data, state)
        state.audit["risk_registry_docx"] = docx_bytes
    except Exception as e:
        state.audit["risk_registry_docx_error"] = str(e)

    return DocumentArtifact(
        doc_type="risk_registry",
        title=doc_title,
        content_markdown=md,
        status="NOT_SUFFICIENT",
        reasons=["Generated by risk_registry_generator; pending validation"],
    )


# =============================================================================
# SECTION 7 — REPAIR ENTRY POINT
# =============================================================================

def repair_risk_registry_artifact(state: PMOState) -> DocumentArtifact:
    """
    Re-runs the full pipeline, injecting validator findings into the prompt
    so the LLM addresses specific gaps on the second pass.
    """
    issues = []
    if "risk_registry" in state.docs:
        issues = state.docs["risk_registry"].reasons or []

    if issues:
        note = ("\n\nREPAIR PASS — the previous generation failed these validation checks. "
                "Ensure the regenerated document addresses ALL of the following:\n" +
                "\n".join(f"  • {i}" for i in issues))
        state.audit.setdefault("uploaded_mapping", {})
        current = state.audit["uploaded_mapping"].get("risk_registry","")
        state.audit["uploaded_mapping"]["risk_registry"] = current + note

    return generate_risk_registry_artifact(state)