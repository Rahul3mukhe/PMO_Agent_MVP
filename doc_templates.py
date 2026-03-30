# doc_templates.py
# Generates Word documents from markdown content.
# Output is styled for formal internal use — suitable for audit, regulatory review, and executive presentation.
 
import io
import os
from datetime import datetime
from typing import List
 
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
 
# ─────────────────────────────────────────────────────────────────────────────
# COLOUR PALETTE
# ─────────────────────────────────────────────────────────────────────────────
_NAVY    = RGBColor(0x00, 0x33, 0x66)   # NTT DATA Darker Blue
_BLUE    = RGBColor(0x00, 0x72, 0xBB)   # NTT DATA Primary Blue
_BODY    = RGBColor(0x1A, 0x1A, 0x2E)
_MID     = RGBColor(0x4A, 0x55, 0x68)
_MUTED   = RGBColor(0x6B, 0x72, 0x80)
_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
_GREEN   = RGBColor(0x15, 0x80, 0x3D)
_RED     = RGBColor(0xB9, 0x1C, 0x1C)
_ROW_A   = "F0F7FF"   # lighter blue for rows
_HDR_BG  = "0072BB"   # table header background (Primary Blue)
 
 
# ─────────────────────────────────────────────────────────────────────────────
# LOW-LEVEL XML HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def _cell_bg(cell, hex_color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)
 
 
def _table_borders(table, color: str = "D1D9E6") -> None:
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcB  = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"),   "single")
                b.set(qn("w:sz"),    "4")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), color)
                tcB.append(b)
            tcPr.append(tcB)
 
 
def _set_page_footer(doc, project_id: str, doc_title: str, org_name: str) -> None:
    section = doc.sections[0]
    footer  = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"{org_name}  |  {doc_title}  |  Ref: {project_id}  |  Page ")
    run.font.size      = Pt(8)
    run.font.color.rgb = _MUTED
    # Auto page number field
    for tag, text in (("begin", None), ("instrText", "PAGE"), ("end", None)):
        if tag == "instrText":
            el = OxmlElement("w:instrText")
            el.text = text
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
        run2 = fp.add_run()
        run2.font.size      = Pt(8)
        run2.font.color.rgb = _MUTED
        run2._r.append(el)
 
 
def _doc_defaults(doc) -> None:
    style      = doc.styles["Normal"]
    style.font.name      = "Calibri"
    style.font.size      = Pt(10)
    style.font.color.rgb = _BODY
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.54)
    sec.right_margin  = Cm(2.54)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
 
 
# ─────────────────────────────────────────────────────────────────────────────
# TYPOGRAPHY HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def _heading(doc, text: str) -> None:
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold           = True
    run.font.size      = Pt(11)
    run.font.color.rgb = _BLUE
    # Blue underline rule
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1A56DB")
    pBdr.append(bot)
    pPr.append(pBdr)
 
 
def _body(doc, text: str, italic: bool = False, color: RGBColor = None) -> None:
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.font.size      = Pt(10)
    run.font.color.rgb = color or _BODY
    run.italic         = italic
 
 
def _bullet(doc, text: str) -> None:
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size      = Pt(10)
    run.font.color.rgb = _BODY
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
 
 
def _cover_bar(doc, org_name: str) -> None:
    """Navy top bar spanning full width."""
    bar  = doc.add_table(rows=1, cols=2)
    bar.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # ── Left: Logo
    cell_logo = bar.cell(0, 0)
    _cell_bg(cell_logo, _HDR_BG)
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "nttdata_logo.png")
    if os.path.exists(logo_path):
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Cm(3.5))
    elif os.path.exists("nttdata_logo.png"):
        run_logo = p_logo.add_run()
        run_logo.add_picture("nttdata_logo.png", width=Cm(3.5))

    # ── Right: Text
    cell_text = bar.cell(0, 1)
    _cell_bg(cell_text, _HDR_BG)
    p_text   = cell_text.paragraphs[0]
    p_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_text.paragraph_format.space_before = Pt(8)
    p_text.paragraph_format.space_after  = Pt(8)
    run = p_text.add_run(f"   {org_name.upper()}   |   GOVERNANCE DOCUMENT")
    run.bold           = True
    run.font.color.rgb = _WHITE
    run.font.size      = Pt(9)
    run.font.name      = "Calibri"
 
 
def _meta_table(doc, rows: List) -> None:
    """Small key-value metadata table."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        tbl.rows[i].cells[0].text = k
        r0 = tbl.rows[i].cells[0].paragraphs[0].runs[0]
        r0.bold           = True
        r0.font.size      = Pt(9)
        r0.font.color.rgb = _MID
        _cell_bg(tbl.rows[i].cells[0], "F0F4FF")
 
        tbl.rows[i].cells[1].text = str(v or "—")
        r1 = tbl.rows[i].cells[1].paragraphs[0].runs[0]
        r1.font.size      = Pt(9)
        r1.font.color.rgb = _BODY
 
    _table_borders(tbl)
    for row in tbl.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11)
    doc.add_paragraph()
 
 
def _sign_off(doc) -> None:
    tbl = doc.add_table(rows=3, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, label in enumerate(("Role", "Name", "Signature / Date")):
        hdr[i].text = label
        run = hdr[i].paragraphs[0].runs[0]
        run.bold           = True
        run.font.size      = Pt(9)
        run.font.color.rgb = _WHITE
        _cell_bg(hdr[i], _HDR_BG)
    for i, role in enumerate(("Project Sponsor", "PMO Representative")):
        tbl.rows[i + 1].cells[0].text = role
        r = tbl.rows[i + 1].cells[0].paragraphs[0].runs[0]
        r.bold      = True
        r.font.size = Pt(9)
    _table_borders(tbl)
    doc.add_paragraph()
 
 
# ─────────────────────────────────────────────────────────────────────────────
# MARKDOWN → WORD BODY PARSER
# Converts the LLM's markdown output into styled Word paragraphs.
# ─────────────────────────────────────────────────────────────────────────────
def _md_to_doc(doc, content: str) -> None:
    """Parse markdown and write into doc using styled paragraphs."""
    for raw in content.splitlines():
        line = raw.rstrip()
 
        # Strip fallback comments injected by llm_providers
        if line.startswith("<!--") and "FALLBACK" in line:
            continue
 
        if line.startswith("### "):
            p   = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(line[4:].strip())
            run.bold           = True
            run.font.size      = Pt(10)
            run.font.color.rgb = _NAVY
 
        elif line.startswith("## "):
            _heading(doc, line[3:].strip())
 
        elif line.startswith("# "):
            _heading(doc, line[2:].strip())
 
        elif line.startswith(("- ", "* ", "+ ")):
            # Strip inline bold markers
            text = line[2:].replace("**", "").replace("__", "").strip()
            _bullet(doc, text)
 
        elif line.strip() == "":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
 
        else:
            # Strip inline bold markers for body text
            text = line.replace("**", "").replace("__", "").strip()
            if text:
                _body(doc, text)
 
 
# ─────────────────────────────────────────────────────────────────────────────
# PUBLIC API — called by pmo_graph.py and app.py
# ─────────────────────────────────────────────────────────────────────────────
 
def create_official_docx(
    org_name: str,
    header: str,
    footer: str,
    title: str,
    md_content: str,
) -> bytes:
    """
    Convert a markdown document into a formatted Word file.
    Returns bytes suitable for st.download_button or file write.
    """
    doc = Document()
    _doc_defaults(doc)
 
    # ── Cover bar
    _cover_bar(doc, org_name)
    doc.add_paragraph()
 
    # ── Document title
    t   = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r   = t.add_run(title)
    r.bold           = True
    r.font.size      = Pt(20)
    r.font.color.rgb = _NAVY
    r.font.name      = "Calibri"
 
    # ── Classification label (header text)
    if header:
        cl = doc.add_paragraph()
        cl.add_run(header).font.size = Pt(9)
        cl.runs[0].italic          = True
        cl.runs[0].font.color.rgb  = _MUTED
 
    doc.add_paragraph()
 
    # ── Metadata row
    _meta_table(doc, [
        ("Organisation",    org_name),
        ("Document",        title),
        ("Classification",  header or "Internal"),
        ("Date",            datetime.now().strftime("%d %B %Y")),
    ])
 
    doc.add_page_break()
 
    # ── Body from markdown
    _md_to_doc(doc, md_content or "")
 
    # ── Sign-off
    _heading(doc, "Approvals")
    _sign_off(doc)
 
    # ── Footer
    _set_page_footer(doc, "—", title, org_name)
 
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
 
 
def create_decision_report_docx(
    org_name: str,
    header: str,
    footer: str,
    project_id: str,
    project_name: str,
    decision: str,
    summary: str,
    gates: list,
    docs: dict,
) -> bytes:
    """
    Generate the PMO Gate Decision Report as a Word file.
    Returns bytes suitable for st.download_button or file write.
    """
    doc = Document()
    _doc_defaults(doc)
 
    # ── Cover bar
    _cover_bar(doc, org_name)
    doc.add_paragraph()
 
    # ── Title
    t   = doc.add_paragraph()
    r   = t.add_run("PMO Gate Decision Report")
    r.bold           = True
    r.font.size      = Pt(20)
    r.font.color.rgb = _NAVY
    r.font.name      = "Calibri"
 
    if header:
        cl = doc.add_paragraph()
        cl.add_run(header).font.size = Pt(9)
        cl.runs[0].italic         = True
        cl.runs[0].font.color.rgb = _MUTED
 
    doc.add_paragraph()
 
    # ── Project metadata
    _meta_table(doc, [
        ("Organisation",    org_name),
        ("Project ID",      project_id),
        ("Project Name",    project_name),
        ("Report Date",     datetime.now().strftime("%d %B %Y %H:%M")),
        ("Classification",  header or "Internal"),
    ])
 
    doc.add_page_break()
 
    # ── 1. Decision
    _heading(doc, "1.  Decision")
    dec_upper = str(decision).strip().upper()
    if "APPROV" in dec_upper or "PASS" in dec_upper:
        dec_color = _GREEN
    elif "REVIEW" in dec_upper or "PENDING" in dec_upper:
        dec_color = RGBColor(0xB4, 0x53, 0x09)
    else:
        dec_color = _RED
    _body(doc, dec_upper, color=dec_color)
    doc.add_paragraph()
 
    # ── 2. Summary
    _heading(doc, "2.  Summary")
    _body(doc, summary or "No summary provided.")
 
    # ── 3. Gate Results
    _heading(doc, "3.  Gate Results")
    gate_items = list(gates.items()) if isinstance(gates, dict) else [(g.gate, g) for g in (gates or [])]
 
    if gate_items:
        tbl = doc.add_table(rows=1 + len(gate_items), cols=3)
        tbl.style = "Table Grid"
        for i, h in enumerate(("Gate", "Result", "Findings")):
            hdr_cell = tbl.rows[0].cells[i]
            hdr_cell.text = h
            run = hdr_cell.paragraphs[0].runs[0]
            run.bold           = True
            run.font.size      = Pt(9)
            run.font.color.rgb = _WHITE
            _cell_bg(hdr_cell, _HDR_BG)
 
        for idx, (gate_name, gate_val) in enumerate(gate_items):
            if isinstance(gate_val, dict):
                passed   = gate_val.get("passed", True)
                findings = gate_val.get("findings", [])
            else:
                passed   = getattr(gate_val, "passed", True)
                findings = getattr(gate_val, "findings", [])
 
            row = tbl.rows[idx + 1].cells
            row[0].text = str(gate_name).replace("_", " ").upper()
            row[1].text = "PASS" if passed else "FAIL"
            row[2].text = "; ".join(str(f) for f in findings) if findings else "—"
 
            bg = "F0FDF4" if passed else "FEF2F2"
            for cell in row:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                _cell_bg(cell, bg)
 
        _table_borders(tbl)
        doc.add_paragraph()
    else:
        _body(doc, "No gate data recorded.", italic=True)
 
    # ── 4. Document Status
    _heading(doc, "4.  Document Status")
    if docs:
        doc_items = list(docs.items()) if isinstance(docs, dict) else []
        if doc_items:
            tbl2 = doc.add_table(rows=1 + len(doc_items), cols=3)
            tbl2.style = "Table Grid"
            for i, h in enumerate(("Document", "Status", "Notes")):
                hdr_cell = tbl2.rows[0].cells[i]
                hdr_cell.text = h
                run = hdr_cell.paragraphs[0].runs[0]
                run.bold           = True
                run.font.size      = Pt(9)
                run.font.color.rgb = _WHITE
                _cell_bg(hdr_cell, _HDR_BG)
 
            for idx, (doc_key, d) in enumerate(doc_items):
                status  = (getattr(d, "status",  None) or (d.get("status")  if isinstance(d, dict) else None) or "—")
                reasons = (getattr(d, "reasons", None) or (d.get("reasons") if isinstance(d, dict) else None) or [])
                title_v = (getattr(d, "title",   None) or (d.get("title")   if isinstance(d, dict) else None) or doc_key)
 
                row = tbl2.rows[idx + 1].cells
                row[0].text = str(title_v)
                row[1].text = str(status)
                row[2].text = "; ".join(str(r) for r in reasons[:2]) if reasons else "—"
 
                bg = "F0FDF4" if status == "SUFFICIENT" else "FEF2F2"
                for cell in row:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    _cell_bg(cell, bg)
 
            _table_borders(tbl2)
            doc.add_paragraph()
 
    # ── 5. Sign-off
    _heading(doc, "5.  Sign-off")
    _sign_off(doc)
 
    # ── Footer
    _set_page_footer(doc, project_id, "PMO Gate Decision Report", org_name)
 
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()