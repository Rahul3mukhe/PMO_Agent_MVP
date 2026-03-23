# risk_registry_generator.py
#
# Drop-in replacement for the risk_registry document type generation.
# Integrates with the existing system:
#   - llm_providers.generate_text()   → structured JSON risk data from LLM
#   - doc_templates                   → Word rendering (python-docx)
#   - guardrails.validate_doc()       → unchanged, validates final markdown
#   - schemas.PMOState / DocumentArtifact → unchanged
#
# HOW IT PLUGS IN:
#   In nodes/generator.py, inside GeneratorNode.__call__ and RepairNode.__call__,
#   replace the generic generate_text call for doc_type == "risk_registry" with:
#
#       from risk_registry_generator import generate_risk_registry_artifact
#       if d == "risk_registry":
#           art = generate_risk_registry_artifact(state)
#           state.docs[d] = art
#           continue
#
#   The function returns a fully populated DocumentArtifact with:
#       - content_markdown : validator-ready markdown (satisfies guardrails)
#       - _docx_bytes      : stored in art.file_path as a sentinel key (see below)
#   The actual docx bytes are also stored in state.audit["risk_registry_docx"]
#   so server.py /export/docx can return them directly without re-rendering.

from __future__ import annotations

import io
import json
import os
import re
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

# ── python-docx ──────────────────────────────────────────────────────────────
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── project imports ───────────────────────────────────────────────────────────
from llm_providers import generate_text
from schemas import DocumentArtifact, PMOState

# =============================================================================
# COLOUR PALETTE  (matches doc_templates.py exactly)
# =============================================================================
_NAVY   = RGBColor(0x00, 0x33, 0x66)
_BLUE   = RGBColor(0x00, 0x72, 0xBB)
_BODY   = RGBColor(0x1A, 0x1A, 0x2E)
_MID    = RGBColor(0x4A, 0x55, 0x68)
_MUTED  = RGBColor(0x6B, 0x72, 0x80)
_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
_GREEN  = RGBColor(0x15, 0x80, 0x3D)
_AMBER  = RGBColor(0xB4, 0x53, 0x09)
_RED    = RGBColor(0xB9, 0x1C, 0x1C)
_ROW_A  = "F0F7FF"
_ROW_B  = "FFFFFF"
_HDR_BG = "003366"

# Hex strings for shading (no RGBColor needed)
_RED_BG    = "FEE2E2"
_AMBER_BG  = "FEF3C7"
_GREEN_BG  = "F0FDF4"
_CRIT_HDR  = "7F1D1D"
_HIGH_HDR  = "78350F"


# =============================================================================
# SECTION 1 — LLM PROMPT
# =============================================================================

def build_risk_llm_prompt(state: PMOState) -> str:
    proj    = state.project
    org     = state.standards["org"]["name"]
    doc_std = state.standards["docs"]["risk_registry"]

    # Collect any uploaded content to give the LLM additional context
    uploaded_md = ""
    if "uploaded_mapping" in state.audit:
        uploaded_md = state.audit["uploaded_mapping"].get("risk_registry", "")

    known_risks_str = "\n".join(f"  - {r}" for r in (proj.known_risks or []))
    deliverables_str = "\n".join(f"  - {d}" for d in (proj.key_deliverables or []))

    return f"""You are a senior PMO Governance Risk Specialist for {org}.
Your task is to produce a COMPLETE, PROFESSIONAL Risk Registry and Risk Assessment
for the project described below. Output ONLY valid JSON — no markdown fences, no
preamble, no explanation.

======= PROJECT CONTEXT =======
Project ID       : {proj.project_id}
Project Name     : {proj.project_name}
Project Type     : {proj.project_type}
Sponsor          : {proj.sponsor or "To be confirmed"}
Est. Budget      : {proj.estimated_budget or "To be confirmed"}
Actual Spend     : {proj.actual_budget_consumed or "Not yet available"}
Timeline (days)  : {proj.total_time_taken_days or "In progress"}
Timeline Summary : {proj.timeline_summary or "Not provided"}
Scope Summary    : {proj.scope_summary or "Not provided"}
Key Deliverables :
{deliverables_str or "  (none listed)"}
Known Risks from intake :
{known_risks_str or "  (none listed — identify risks from context)"}

======= UPLOADED DOCUMENT EXTRACT =======
{uploaded_md[:4000] if uploaded_md else "(no uploaded document — derive risks from project context)"}

======= OUTPUT SCHEMA =======
Return a single JSON object with this EXACT structure:

{{
  "executive_summary": "2-3 paragraph professional summary of the programme risk posture",
  "risk_framework": "1-2 paragraphs describing the risk methodology used (ISO 31000, likelihood×impact 5×5 matrix)",
  "assumptions": [
    "Assumption statement 1",
    "Assumption statement 2"
    // minimum 6 assumptions
  ],
  "risks": [
    {{
      "id": "R-001",
      "category": "Technical | Resource | Regulatory | Schedule | Financial | Security | Organisational",
      "title": "Short risk title (max 8 words)",
      "description": "Full description of the risk, root cause, and why it matters for this project",
      "likelihood": "Low | Medium | High",
      "impact": "Low | Medium | High | Critical",
      "rating": "Low | Medium | High | Critical",
      "score": 12,
      "owner": "Role title (e.g. Head of Architecture)",
      "status": "Open | Under Review | Active — Mitigating | Closed",
      "mitigation": "Specific, actionable mitigation steps for this project",
      "contingency": "Specific contingency plan if the risk materialises",
      "assumption": "Key assumption underlying this risk assessment",
      "review_date": "Weekly | Fortnightly | Monthly",
      "residual_rating": "Low | Medium | High"
    }}
    // minimum 6 risks — more if project context warrants it
  ],
  "review_cadence": [
    {{
      "review_type": "e.g. Weekly Risk Standup",
      "frequency": "e.g. Every Monday 09:00",
      "attendees": "e.g. Risk Owners + Delivery Leads",
      "scope": "e.g. Critical & High risks",
      "output": "e.g. Risk Action Log"
    }}
    // minimum 4 review cadence entries
  ]
}}

CRITICAL RULES:
1. Every field in every object MUST be populated — never null, never empty string.
2. Risk scores: Low=1-3, Medium=4-8, High=9-15, Critical=16-25 (likelihood×impact on 5×5 scale).
3. All mitigations and contingencies MUST be specific to THIS project — not generic.
4. Owner field: always a role title, never a person's name.
5. Output ONLY the JSON object. Nothing else.
""".strip()


# =============================================================================
# SECTION 2 — LLM CALL + JSON PARSE
# =============================================================================

def _parse_risk_json(raw: str) -> Dict[str, Any]:
    """Extract and parse JSON from LLM output robustly."""
    # Strip markdown fences if present
    text = re.sub(r"```(?:json)?", "", raw).strip().rstrip("`").strip()
    # Find outermost braces
    start = text.find("{")
    end   = text.rfind("}") + 1
    if start == -1 or end == 0:
        raise ValueError("No JSON object found in LLM response")
    return json.loads(text[start:end])


def _validate_and_fill_risk_data(data: Dict[str, Any], proj) -> Dict[str, Any]:
    """Ensure all required fields are present; fill defaults for anything missing."""
    defaults = {
        "executive_summary": f"Risk assessment for {proj.project_name}.",
        "risk_framework": "This register follows ISO 31000 risk management guidelines using a 5×5 likelihood-impact matrix.",
        "assumptions": [
            "Project scope remains stable after sign-off.",
            "Named resources remain available throughout the project lifecycle.",
            "Stakeholder approvals will be granted within agreed SLA timeframes.",
            "No material changes to applicable regulatory requirements during delivery.",
            "Budget contingency reserve is accessible upon PMO approval.",
            "Third-party vendors will honour contractual SLA commitments.",
        ],
        "risks": [],
        "review_cadence": [
            {"review_type": "Weekly Risk Standup", "frequency": "Weekly",
             "attendees": "Risk Owners + PM", "scope": "High & Critical risks",
             "output": "Risk Action Log"},
            {"review_type": "Programme Risk Review", "frequency": "Fortnightly",
             "attendees": "Programme Manager + Risk Owners", "scope": "All active risks",
             "output": "Risk Register Update"},
            {"review_type": "Steering Committee Report", "frequency": "Monthly",
             "attendees": "Sponsor + SteerCo", "scope": "Escalations",
             "output": "SteerCo Minutes"},
            {"review_type": "Full Register Audit", "frequency": "Quarterly",
             "attendees": "PMO + Internal Audit", "scope": "Full register",
             "output": "Audit Report"},
        ],
    }
    for k, v in defaults.items():
        if k not in data or not data[k]:
            data[k] = v

    # Validate each risk entry
    risk_defaults = {
        "id": "R-000", "category": "General", "title": "Unnamed Risk",
        "description": "Risk details to be documented.", "likelihood": "Medium",
        "impact": "Medium", "rating": "Medium", "score": 9,
        "owner": "Programme Manager", "status": "Open",
        "mitigation": "Mitigation to be defined.", "contingency": "Contingency to be defined.",
        "assumption": "Standard project assumptions apply.",
        "review_date": "Monthly", "residual_rating": "Low",
    }
    for i, r in enumerate(data.get("risks", [])):
        for k, v in risk_defaults.items():
            if k not in r or not r[k]:
                r[k] = v
        # Auto-generate ID if missing
        if r["id"] == "R-000":
            r["id"] = f"R-{str(i+1).zfill(3)}"

    return data


def call_llm_for_risk_data(state: PMOState) -> Dict[str, Any]:
    """Call the LLM and return parsed, validated risk data dict."""
    prompt = build_risk_llm_prompt(state)

    raw = generate_text(
        provider=state.provider,
        model=state.model,
        prompt=prompt,
        project=state.project,
        doc_type="risk_registry",
        standards=state.standards,
        api_key=state.audit.get("api_key"),
        temperature=0.1,          # slight warmth for richer text; 0 for strict JSON
        max_tokens=state.audit.get("max_tokens", 4096),
    )

    try:
        data = _parse_risk_json(raw)
    except (ValueError, json.JSONDecodeError) as e:
        # LLM returned non-JSON — fall back to local template enriched with
        # whatever the LLM did produce as a prose note
        state.audit["risk_registry_parse_error"] = f"{e}: {raw[:200]}"
        data = {}

    return _validate_and_fill_risk_data(data, state.project)


# =============================================================================
# SECTION 3 — MARKDOWN GENERATOR  (for guardrails validation)
# =============================================================================
# The guardrails.validate_doc() function checks the markdown content of the
# DocumentArtifact.  We generate standards-compliant markdown here so
# validate_doc passes without needing a repair cycle.

def _risk_rating_line(r: dict) -> str:
    return f"Likelihood: {r['likelihood']} | Impact: {r['impact']} | Score: {r['score']} | Rating: {r['rating']}"


def build_risk_markdown(data: Dict[str, Any], proj, standards: dict) -> str:
    """
    Build markdown that satisfies all guardrails checks for risk_registry:
      - required_sections present as ## headings
      - min_total_lines >= 20
      - min_bullets >= 5
      - no rejected keywords
    """
    doc_std = standards["docs"]["risk_registry"]
    required_sections = doc_std.get("required_sections", [])

    lines: List[str] = []

    def h(title: str):
        lines.append(f"## {title}")
        lines.append("")

    def body(text: str):
        lines.append(text)
        lines.append("")

    def bullet(text: str):
        lines.append(f"- {text}")

    risks = data.get("risks", [])

    # ── Map each required section to content ──────────────────────────────────
    # We normalise to lowercase for matching so the YAML section names are
    # handled flexibly (e.g. "Risk Summary" → "risk summary")
    section_lower_map = {s.lower(): s for s in required_sections}

    def emit_section(key_lower: str, original: str):
        h(original)
        sl = key_lower

        if sl == "overview":
            body(f"This Risk Registry and Risk Assessment documents all identified risks for "
                 f"{proj.project_name} (Ref: {proj.project_id}). It is maintained in accordance "
                 f"with {standards['org']['name']} PMO governance standards and must be reviewed "
                 f"and approved before each project gate.")
            body(data.get("risk_framework", ""))

        elif sl == "risk summary":
            # Rating summary bullets — counts by level
            rating_counts: Dict[str, int] = {}
            for r in risks:
                rating_counts[r.get("rating", "Unknown")] = rating_counts.get(r.get("rating", "Unknown"), 0) + 1
            body(data.get("executive_summary", ""))
            for rating, count in rating_counts.items():
                bullet(f"{count} risk(s) rated {rating}")
            lines.append("")
            bullet(f"Total risks identified: {len(risks)}")
            lines.append("")

        elif sl == "detailed risks":
            for r in risks:
                bullet(f"{r['id']}: {r['title']} — {r['description'][:120]}")
            lines.append("")

        elif sl == "mitigations":
            for r in risks:
                bullet(f"{r['id']} ({r['title']}): {r['mitigation']}")
            lines.append("")

        elif sl == "owners":
            seen_owners: Dict[str, List[str]] = {}
            for r in risks:
                seen_owners.setdefault(r["owner"], []).append(r["id"])
            for owner, ids in seen_owners.items():
                bullet(f"{owner} — responsible for: {', '.join(ids)}")
            lines.append("")
            body("All risk owners are accountable for monitoring assigned risks and escalating "
                 "status changes to the Programme Manager within 48 hours of any material change.")

        elif sl == "registry overview":
            body(f"This registry covers {len(risks)} risks across "
                 f"{len(set(r['category'] for r in risks))} categories. "
                 f"Risk ratings follow a 5×5 likelihood-impact matrix aligned to ISO 31000.")
            body("The register is version-controlled and retained for a minimum of 7 years "
                 "in accordance with document retention policy.")

        elif sl == "risk list":
            for r in risks:
                bullet(f"{r['id']} | {r['title']} | {r['category']} | "
                       f"{_risk_rating_line(r)} | Owner: {r['owner']} | Status: {r['status']}")
            lines.append("")

        elif sl == "review cadence":
            cadence = data.get("review_cadence", [])
            for entry in cadence:
                bullet(f"{entry.get('review_type', '—')}: {entry.get('frequency', '—')} — "
                       f"Attendees: {entry.get('attendees', '—')} — Output: {entry.get('output', '—')}")
            lines.append("")

        elif sl == "approvals":
            body("The following roles are required to review and approve this document:")
            bullet("Project Sponsor — Name: _________________ | Signature: _____________ | Date: _______")
            bullet("Programme Manager — Name: _________________ | Signature: _____________ | Date: _______")
            bullet("PMO Representative — Name: _________________ | Signature: _____________ | Date: _______")
            bullet("Chief Risk Officer — Name: _________________ | Signature: _____________ | Date: _______")
            lines.append("")

        else:
            # Generic fallback for any section in the YAML not explicitly handled
            body(f"Details for section '{original}' are documented below.")
            for r in risks[:5]:
                bullet(f"{r['id']}: {r['title']}")
            lines.append("")

    # Emit every required section in order
    for section in required_sections:
        emit_section(section.lower(), section)

    # Assumptions — always include even if not in required_sections
    if "assumptions" not in [s.lower() for s in required_sections]:
        h("Assumptions")
    else:
        h("Assumptions")
    for a in data.get("assumptions", []):
        bullet(a)
    lines.append("")

    return "\n".join(lines)


# =============================================================================
# SECTION 4 — WORD DOCUMENT RENDERER  (python-docx, mirrors doc_templates.py)
# =============================================================================

# ── XML / style helpers ───────────────────────────────────────────────────────

def _cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _table_borders(table, color: str = "D1D9E6") -> None:
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcB  = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"),   "single")
                b.set(qn("w:sz"),    "4")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), color)
                tcB.append(b)
            tcPr.append(tcB)


def _set_col_width(cell, width_cm: float) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(width_cm * 567)))  # 1 cm ≈ 567 twips
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _doc_defaults(doc: Document) -> None:
    style           = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = _BODY
    sec               = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(1.8)
    sec.right_margin  = Cm(1.8)
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)


def _page_footer(doc: Document, project_id: str, org_name: str) -> None:
    section = doc.sections[0]
    footer  = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(f"{org_name}  |  Risk Registry & Risk Assessment  |  {project_id}  |  Page ")
    run.font.size      = Pt(8)
    run.font.color.rgb = _MUTED
    for tag, text in (("begin", None), ("instrText", "PAGE"), ("end", None)):
        if tag == "instrText":
            el = OxmlElement("w:instrText")
            el.text = text
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
        r2 = fp.add_run()
        r2.font.size      = Pt(8)
        r2.font.color.rgb = _MUTED
        r2._r.append(el)


# ── Typography helpers ────────────────────────────────────────────────────────

def _section_heading(doc: Document, text: str) -> None:
    """Blue ALL-CAPS section heading with underline rule — matches doc_templates._heading()."""
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    run.bold           = True
    run.font.size      = Pt(12)
    run.font.color.rgb = _BLUE
    run.font.name      = "Calibri"
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "0072BB")
    pBdr.append(bot)
    pPr.append(pBdr)


def _sub_heading(doc: Document, text: str) -> None:
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(11)
    run.font.color.rgb = _NAVY
    run.font.name      = "Calibri"


def _body_para(doc: Document, text: str, italic: bool = False,
               color: Optional[RGBColor] = None) -> None:
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.font.size      = Pt(10)
    run.font.color.rgb = color or _BODY
    run.italic         = italic
    run.font.name      = "Calibri"


def _bullet(doc: Document, text: str) -> None:
    p   = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size      = Pt(10)
    run.font.color.rgb = _BODY
    run.font.name      = "Calibri"
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)


def _gap(doc: Document, space_pt: float = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_pt)
    p.paragraph_format.space_after  = Pt(0)


# ── Risk colour helpers ───────────────────────────────────────────────────────

def _risk_rgb(level: str) -> RGBColor:
    l = level.upper()
    if l in ("CRITICAL", "HIGH"):
        return _RED
    if l == "MEDIUM":
        return _AMBER
    return _GREEN


def _risk_bg_hex(level: str) -> str:
    l = level.upper()
    if l == "CRITICAL":
        return _RED_BG
    if l == "HIGH":
        return "FEF3F2"
    if l == "MEDIUM":
        return _AMBER_BG
    return _GREEN_BG


# ── Cover page ────────────────────────────────────────────────────────────────

def _cover_page(doc: Document, proj, standards: dict, logo_path: Optional[str]) -> None:
    org    = standards["org"]["name"]
    header = standards["org"].get("doc_header", "Internal")

    # ── Top banner bar (2-column table, full width) ───────────────────────────
    bar = doc.add_table(rows=1, cols=2)
    bar.alignment = WD_TABLE_ALIGNMENT.LEFT

    logo_cell = bar.cell(0, 0)
    _cell_bg(logo_cell, _HDR_BG)
    p_logo = logo_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_before = Pt(6)
    p_logo.paragraph_format.space_after  = Pt(6)
    if logo_path and os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Cm(3.5))
    else:
        run = p_logo.add_run(org.upper())
        run.bold           = True
        run.font.size      = Pt(14)
        run.font.color.rgb = _WHITE
        run.font.name      = "Calibri"

    text_cell = bar.cell(0, 1)
    _cell_bg(text_cell, _HDR_BG)
    p_text = text_cell.paragraphs[0]
    p_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_text.paragraph_format.space_before = Pt(10)
    p_text.paragraph_format.space_after  = Pt(10)
    run = p_text.add_run(f"   {org.upper()}   |   GOVERNANCE DOCUMENT   ")
    run.bold           = True
    run.font.color.rgb = _WHITE
    run.font.size      = Pt(9)
    run.font.name      = "Calibri"

    _set_col_width(logo_cell, 4.0)
    _set_col_width(text_cell, 13.2)

    doc.add_paragraph()

    # ── Document type label ───────────────────────────────────────────────────
    lbl = doc.add_paragraph()
    lbl_run = lbl.add_run("PROGRAMME RISK DOCUMENTATION")
    lbl_run.font.size      = Pt(9)
    lbl_run.font.color.rgb = _MUTED
    lbl_run.font.name      = "Calibri"
    lbl_run.bold           = True

    # ── Main title ────────────────────────────────────────────────────────────
    t1 = doc.add_paragraph()
    r1 = t1.add_run("Risk Registry")
    r1.bold           = True
    r1.font.size      = Pt(32)
    r1.font.color.rgb = _NAVY
    r1.font.name      = "Calibri"

    t2 = doc.add_paragraph()
    r2 = t2.add_run("& Risk Assessment")
    r2.bold           = True
    r2.font.size      = Pt(32)
    r2.font.color.rgb = _BLUE
    r2.font.name      = "Calibri"
    t2.paragraph_format.space_after = Pt(6)

    # ── Subtitle: project name ────────────────────────────────────────────────
    sub = doc.add_paragraph()
    sub_run = sub.add_run(proj.project_name)
    sub_run.bold           = True
    sub_run.font.size      = Pt(14)
    sub_run.font.color.rgb = _MID
    sub_run.font.name      = "Calibri"

    doc.add_paragraph()

    # ── Metadata table ────────────────────────────────────────────────────────
    meta_rows = [
        ("Project Reference", proj.project_id),
        ("Project Type",      proj.project_type),
        ("Programme Sponsor", proj.sponsor or "To be confirmed"),
        ("Organisation",      org),
        ("Classification",    header),
        ("Document Version",  "v1.0 — Initial Release"),
        ("Prepared By",       f"{org} PMO Risk Team"),
        ("Date",              datetime.now().strftime("%d %B %Y")),
    ]
    tbl = doc.add_table(rows=len(meta_rows), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(meta_rows):
        tbl.rows[i].cells[0].text = k
        r0 = tbl.rows[i].cells[0].paragraphs[0].runs[0]
        r0.bold           = True
        r0.font.size      = Pt(9)
        r0.font.color.rgb = _WHITE
        _cell_bg(tbl.rows[i].cells[0], _HDR_BG)
        tbl.rows[i].cells[1].text = str(v or "—")
        r1 = tbl.rows[i].cells[1].paragraphs[0].runs[0]
        r1.font.size      = Pt(9)
        r1.font.color.rgb = _BODY
        _cell_bg(tbl.rows[i].cells[1], _ROW_A)
    _table_borders(tbl)
    for row in tbl.rows:
        _set_col_width(row.cells[0], 4.5)
        _set_col_width(row.cells[1], 11.0)

    doc.add_paragraph()

    # ── Confidentiality notice ────────────────────────────────────────────────
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf.add_run(
        "COMMERCIAL-IN-CONFIDENCE  |  FOR INTERNAL USE ONLY  |  DO NOT DISTRIBUTE"
    )
    conf_run.font.size      = Pt(8)
    conf_run.italic         = True
    conf_run.font.color.rgb = _MUTED
    conf_run.font.name      = "Calibri"

    doc.add_page_break()


# ── Section renderers ─────────────────────────────────────────────────────────

def _render_exec_summary(doc: Document, data: dict, risks: list) -> None:
    _section_heading(doc, "1.  Executive Summary")

    for para in data["executive_summary"].split("\n\n"):
        if para.strip():
            _body_para(doc, para.strip())

    _gap(doc)

    # Stats row
    counts = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0}
    for r in risks:
        level = r.get("rating", "Low")
        if level in counts:
            counts[level] += 1

    stat_data = [
        ("Total Risks",    str(len(risks)),   _BODY),
        ("Critical/High",  str(counts["Critical"] + counts["High"]), _RED),
        ("Medium",         str(counts["Medium"]),  _AMBER),
        ("Low",            str(counts["Low"]),     _GREEN),
    ]
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"
    for i, (label, val, color) in enumerate(stat_data):
        tbl.rows[0].cells[i].text = label
        lrun = tbl.rows[0].cells[i].paragraphs[0].runs[0]
        lrun.bold = True; lrun.font.size = Pt(9); lrun.font.color.rgb = _WHITE
        _cell_bg(tbl.rows[0].cells[i], _HDR_BG)

        tbl.rows[1].cells[i].text = val
        vrun = tbl.rows[1].cells[i].paragraphs[0].runs[0]
        vrun.bold = True; vrun.font.size = Pt(22); vrun.font.color.rgb = color
        tbl.rows[1].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _cell_bg(tbl.rows[1].cells[i], _ROW_A)
    _table_borders(tbl)


def _render_framework(doc: Document, data: dict) -> None:
    doc.add_page_break()
    _section_heading(doc, "2.  Risk Framework & Methodology")
    _body_para(doc, data.get("risk_framework", ""))

    _gap(doc)
    _sub_heading(doc, "Risk Scoring Matrix  (Likelihood × Impact)")

    likelihoods = [
        ("Almost Certain (5)", 5),
        ("Likely (4)",         4),
        ("Possible (3)",       3),
        ("Unlikely (2)",       2),
        ("Rare (1)",           1),
    ]
    impacts = [
        ("Negligible (1)", 1),
        ("Minor (2)",      2),
        ("Moderate (3)",   3),
        ("Major (4)",      4),
        ("Catastrophic (5)", 5),
    ]

    tbl = doc.add_table(rows=len(likelihoods) + 1, cols=len(impacts) + 1)
    tbl.style = "Table Grid"

    # Header row
    tbl.rows[0].cells[0].text = "Likelihood \\ Impact"
    hr = tbl.rows[0].cells[0].paragraphs[0].runs[0]
    hr.bold = True; hr.font.size = Pt(8); hr.font.color.rgb = _WHITE
    _cell_bg(tbl.rows[0].cells[0], _HDR_BG)

    for j, (ilabel, _) in enumerate(impacts):
        tbl.rows[0].cells[j + 1].text = ilabel
        r = tbl.rows[0].cells[j + 1].paragraphs[0].runs[0]
        r.bold = True; r.font.size = Pt(8); r.font.color.rgb = _WHITE
        _cell_bg(tbl.rows[0].cells[j + 1], _HDR_BG)

    for i, (llabel, lval) in enumerate(likelihoods):
        tbl.rows[i + 1].cells[0].text = llabel
        lr = tbl.rows[i + 1].cells[0].paragraphs[0].runs[0]
        lr.bold = True; lr.font.size = Pt(8); lr.font.color.rgb = _WHITE
        _cell_bg(tbl.rows[i + 1].cells[0], _HDR_BG)

        for j, (_, ival) in enumerate(impacts):
            score = lval * ival
            cell  = tbl.rows[i + 1].cells[j + 1]
            cell.text = str(score)
            sr = cell.paragraphs[0].runs[0]
            sr.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if score >= 16:
                _cell_bg(cell, _RED_BG);   sr.bold = True;  sr.font.color.rgb = _RED
            elif score >= 9:
                _cell_bg(cell, _AMBER_BG); sr.font.color.rgb = _AMBER
            else:
                _cell_bg(cell, _GREEN_BG); sr.font.color.rgb = _GREEN
    _table_borders(tbl)

    _gap(doc, 10)

    # Legend
    _sub_heading(doc, "Rating Legend")
    legend_rows = [
        ("16–25", "CRITICAL", _RED_BG,   _RED),
        ("9–15",  "HIGH",     _AMBER_BG, _AMBER),
        ("4–8",   "MEDIUM",   _GREEN_BG, _GREEN),
        ("1–3",   "LOW",      _ROW_B,    _GREEN),
    ]
    ltbl = doc.add_table(rows=len(legend_rows) + 1, cols=4)
    ltbl.style = "Table Grid"
    for i, h in enumerate(("Score", "Rating", "Escalation Requirement", "Review Frequency")):
        ltbl.rows[0].cells[i].text = h
        hr2 = ltbl.rows[0].cells[i].paragraphs[0].runs[0]
        hr2.bold = True; hr2.font.size = Pt(9); hr2.font.color.rgb = _WHITE
        _cell_bg(ltbl.rows[0].cells[i], _HDR_BG)

    escalations = [
        "Immediate Steering Committee notification",
        "Programme Sponsor + PM within 24 hours",
        "Programme Manager review",
        "Risk Owner monitoring",
    ]
    freqs = ["Weekly", "Fortnightly", "Monthly", "Monthly"]
    for i, (score, rating, bg, color) in enumerate(legend_rows):
        ltbl.rows[i+1].cells[0].text = score
        ltbl.rows[i+1].cells[1].text = rating
        ltbl.rows[i+1].cells[2].text = escalations[i]
        ltbl.rows[i+1].cells[3].text = freqs[i]
        for ci in range(4):
            r = ltbl.rows[i+1].cells[ci].paragraphs[0].runs[0]
            r.font.size = Pt(9)
            if ci == 1:
                r.bold = True; r.font.color.rgb = color
            _cell_bg(ltbl.rows[i+1].cells[ci], bg)
    _table_borders(ltbl)


def _render_assumptions(doc: Document, data: dict) -> None:
    doc.add_page_break()
    _section_heading(doc, "3.  Risk Assumptions Register")
    _body_para(doc, (
        "The following assumptions underpin the risk assessments documented in this register. "
        "Each assumption represents a condition that, if it does not hold, will trigger "
        "reassessment of associated risks and may require escalation to the Steering Committee."
    ))
    _gap(doc)

    assumptions = data.get("assumptions", [])
    tbl = doc.add_table(rows=len(assumptions) + 1, cols=3)
    tbl.style = "Table Grid"
    for i, h in enumerate(("Ref", "Assumption Statement", "Status")):
        tbl.rows[0].cells[i].text = h
        hr = tbl.rows[0].cells[i].paragraphs[0].runs[0]
        hr.bold = True; hr.font.size = Pt(9); hr.font.color.rgb = _WHITE
        _cell_bg(tbl.rows[0].cells[i], _HDR_BG)

    for i, assumption in enumerate(assumptions):
        bg = _ROW_A if i % 2 == 0 else _ROW_B
        tbl.rows[i+1].cells[0].text = f"A-{str(i+1).zfill(3)}"
        tbl.rows[i+1].cells[1].text = assumption
        tbl.rows[i+1].cells[2].text = "Valid"
        r0 = tbl.rows[i+1].cells[0].paragraphs[0].runs[0]
        r0.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = _BLUE
        _cell_bg(tbl.rows[i+1].cells[0], bg)
        r1 = tbl.rows[i+1].cells[1].paragraphs[0].runs[0]
        r1.font.size = Pt(9); r1.font.color.rgb = _BODY
        _cell_bg(tbl.rows[i+1].cells[1], bg)
        r2 = tbl.rows[i+1].cells[2].paragraphs[0].runs[0]
        r2.bold = True; r2.font.size = Pt(9); r2.font.color.rgb = _GREEN
        _cell_bg(tbl.rows[i+1].cells[2], bg)
    _table_borders(tbl)
    for row in tbl.rows:
        _set_col_width(row.cells[0], 1.5)
        _set_col_width(row.cells[1], 13.5)
        _set_col_width(row.cells[2], 2.0)


def _render_risk_summary_dashboard(doc: Document, risks: list) -> None:
    doc.add_page_break()
    _section_heading(doc, "4.  Risk Summary Dashboard")
    _body_para(doc, (
        "The table below provides a consolidated single-page view of all identified risks, "
        "their current ratings, owners, and status. Full details are in Section 5."
    ))
    _gap(doc)

    cols = ("ID", "Risk Title", "Category", "Likelihood", "Impact", "Score", "Owner", "Status")
    widths = (1.2, 4.5, 2.0, 2.0, 2.0, 1.2, 3.0, 2.6)

    tbl = doc.add_table(rows=len(risks) + 1, cols=len(cols))
    tbl.style = "Table Grid"
    for i, (h, w) in enumerate(zip(cols, widths)):
        tbl.rows[0].cells[i].text = h
        hr = tbl.rows[0].cells[i].paragraphs[0].runs[0]
        hr.bold = True; hr.font.size = Pt(9); hr.font.color.rgb = _WHITE
        _cell_bg(tbl.rows[0].cells[i], _HDR_BG)
        _set_col_width(tbl.rows[0].cells[i], w)

    for i, r in enumerate(risks):
        bg = _ROW_A if i % 2 == 0 else _ROW_B
        vals = [r["id"], r["title"], r["category"], r["likelihood"],
                r["impact"], str(r["score"]), r["owner"], r["status"]]
        for j, (val, w) in enumerate(zip(vals, widths)):
            cell = tbl.rows[i+1].cells[j]
            cell.text = val
            run  = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            _set_col_width(cell, w)
            if j == 0:       # ID
                run.bold = True; run.font.color.rgb = _BLUE; _cell_bg(cell, bg)
            elif j in (3, 4, 5):  # likelihood / impact / score
                _cell_bg(cell, _risk_bg_hex(r["rating"]))
                run.bold = True; run.font.color.rgb = _risk_rgb(r["rating"])
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                run.font.color.rgb = _BODY; _cell_bg(cell, bg)
    _table_borders(tbl)


def _render_detailed_risks(doc: Document, risks: list) -> None:
    doc.add_page_break()
    _section_heading(doc, "5.  Risk Registry — Detailed Entries")
    _body_para(doc, (
        "Each entry below provides the full assessment including description, likelihood and "
        "impact rationale, mitigation actions, contingency plan, assumptions, and residual rating."
    ))

    for idx, r in enumerate(risks):
        if idx > 0:
            _gap(doc, 16)

        # ── Risk header bar ───────────────────────────────────────────────────
        hdr_tbl = doc.add_table(rows=1, cols=3)
        hdr_tbl.style = "Table Grid"
        hdr_tbl.rows[0].cells[0].text = r["id"]
        id_run = hdr_tbl.rows[0].cells[0].paragraphs[0].runs[0]
        id_run.bold = True; id_run.font.size = Pt(11); id_run.font.color.rgb = _WHITE
        _cell_bg(hdr_tbl.rows[0].cells[0], _HDR_BG)
        _set_col_width(hdr_tbl.rows[0].cells[0], 1.5)

        hdr_tbl.rows[0].cells[1].text = f"{r['title']}  [{r['category']}]"
        title_run = hdr_tbl.rows[0].cells[1].paragraphs[0].runs[0]
        title_run.bold = True; title_run.font.size = Pt(11); title_run.font.color.rgb = _WHITE
        _cell_bg(hdr_tbl.rows[0].cells[1], _HDR_BG)
        _set_col_width(hdr_tbl.rows[0].cells[1], 12.5)

        score_text = f"Score: {r['score']}  [{r['rating']}]"
        hdr_tbl.rows[0].cells[2].text = score_text
        score_run = hdr_tbl.rows[0].cells[2].paragraphs[0].runs[0]
        score_run.bold = True; score_run.font.size = Pt(11); score_run.font.color.rgb = _WHITE
        _cell_bg(hdr_tbl.rows[0].cells[2], _HDR_BG)
        _set_col_width(hdr_tbl.rows[0].cells[2], 2.5)
        hdr_tbl.rows[0].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _table_borders(hdr_tbl, "003366")

        # ── Detail rows ───────────────────────────────────────────────────────
        detail_fields = [
            ("Description",       r["description"]),
            ("Likelihood",        f"{r['likelihood']} — assessed probability that this risk materialises within the current project phase."),
            ("Impact",            f"{r['impact']} — consequence on scope, cost, timeline, or regulatory standing if the risk materialises."),
            ("Risk Rating",       f"{r['rating']} (Score: {r['score']})"),
            ("Mitigation Actions",r["mitigation"]),
            ("Contingency Plan",  r["contingency"]),
            ("Key Assumption",    r["assumption"]),
            ("Risk Owner",        r["owner"]),
            ("Review Frequency",  r["review_date"]),
            ("Residual Rating",   f"{r['residual_rating']} (post-mitigation)"),
            ("Current Status",    r["status"]),
        ]

        det_tbl = doc.add_table(rows=len(detail_fields), cols=2)
        det_tbl.style = "Table Grid"
        for i, (key, val) in enumerate(detail_fields):
            bg = _ROW_A if i % 2 == 0 else _ROW_B
            det_tbl.rows[i].cells[0].text = key
            kr = det_tbl.rows[i].cells[0].paragraphs[0].runs[0]
            kr.bold = True; kr.font.size = Pt(9); kr.font.color.rgb = _WHITE
            _cell_bg(det_tbl.rows[i].cells[0], _HDR_BG)
            _set_col_width(det_tbl.rows[i].cells[0], 3.5)

            det_tbl.rows[i].cells[1].text = val
            vr = det_tbl.rows[i].cells[1].paragraphs[0].runs[0]
            vr.font.size = Pt(9)
            if key in ("Risk Rating", "Residual Rating"):
                lvl = r["rating"] if key == "Risk Rating" else r["residual_rating"]
                _cell_bg(det_tbl.rows[i].cells[1], _risk_bg_hex(lvl))
                vr.bold = True; vr.font.color.rgb = _risk_rgb(lvl)
            else:
                _cell_bg(det_tbl.rows[i].cells[1], bg)
                vr.font.color.rgb = _BODY
            _set_col_width(det_tbl.rows[i].cells[1], 13.0)
        _table_borders(det_tbl)


def _render_risk_mitigation_matrix(doc: Document, risks: list) -> None:
    doc.add_page_break()
    _section_heading(doc, "6.  Risk Assessment Matrix — Risk vs. Mitigation")
    _body_para(doc, (
        "The matrix below maps each risk directly to its mitigation strategy and contingency plan, "
        "providing auditors and steering committee members a single-view reference for risk "
        "treatment decisions."
    ))
    _gap(doc)

    cols    = ("ID", "Risk Title", "Rating", "Mitigation Strategy", "Contingency Plan", "Residual")
    widths  = (1.2, 3.0, 1.8, 4.8, 4.0, 1.8)

    tbl = doc.add_table(rows=len(risks) + 1, cols=len(cols))
    tbl.style = "Table Grid"
    for i, (h, w) in enumerate(zip(cols, widths)):
        tbl.rows[0].cells[i].text = h
        hr = tbl.rows[0].cells[i].paragraphs[0].runs[0]
        hr.bold = True; hr.font.size = Pt(9); hr.font.color.rgb = _WHITE
        _cell_bg(tbl.rows[0].cells[i], _HDR_BG)
        _set_col_width(tbl.rows[0].cells[i], w)

    for i, r in enumerate(risks):
        bg = _ROW_A if i % 2 == 0 else _ROW_B
        row_vals = [
            r["id"], r["title"], r["rating"],
            r["mitigation"],
            r["contingency"],
            r["residual_rating"],
        ]
        for j, (val, w) in enumerate(zip(row_vals, widths)):
            cell = tbl.rows[i+1].cells[j]
            cell.text = val
            run  = cell.paragraphs[0].runs[0]
            run.font.size = Pt(8)
            _set_col_width(cell, w)
            if j == 0:
                run.bold = True; run.font.color.rgb = _BLUE; _cell_bg(cell, bg)
            elif j in (2, 5):   # rating / residual
                lvl = r["rating"] if j == 2 else r["residual_rating"]
                _cell_bg(cell, _risk_bg_hex(lvl))
                run.bold = True; run.font.color.rgb = _risk_rgb(lvl)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                run.font.color.rgb = _BODY; _cell_bg(cell, bg)
    _table_borders(tbl)


def _render_risk_owners(doc: Document, risks: list) -> None:
    doc.add_page_break()
    _section_heading(doc, "7.  Risk Owners Register")
    _body_para(doc, (
        "Risk owners are accountable for monitoring assigned risks, maintaining mitigation action "
        "logs, and escalating status changes to the Programme Manager within 48 hours of any "
        "material change in risk exposure."
    ))
    _gap(doc)

    owner_map: Dict[str, List[str]] = {}
    for r in risks:
        owner_map.setdefault(r["owner"], []).append(r["id"])

    tbl = doc.add_table(rows=len(owner_map) + 1, cols=4)
    tbl.style = "Table Grid"
    for i, h in enumerate(("Risk Owner (Role)", "Assigned Risk IDs", "Reports To", "Review Frequency")):
        tbl.rows[0].cells[i].text = h
        hr = tbl.rows[0].cells[i].paragraphs[0].runs[0]
        hr.bold = True; hr.font.size = Pt(9); hr.font.color.rgb = _WHITE
        _cell_bg(tbl.rows[0].cells[i], _HDR_BG)

    for i, (owner, ids) in enumerate(owner_map.items()):
        bg = _ROW_A if i % 2 == 0 else _ROW_B
        tbl.rows[i+1].cells[0].text = owner
        r0 = tbl.rows[i+1].cells[0].paragraphs[0].runs[0]
        r0.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = _BODY
        _cell_bg(tbl.rows[i+1].cells[0], bg)

        tbl.rows[i+1].cells[1].text = ", ".join(ids)
        r1 = tbl.rows[i+1].cells[1].paragraphs[0].runs[0]
        r1.font.size = Pt(9); r1.font.color.rgb = _BLUE; r1.bold = True
        _cell_bg(tbl.rows[i+1].cells[1], bg)

        tbl.rows[i+1].cells[2].text = "Programme Manager"
        r2 = tbl.rows[i+1].cells[2].paragraphs[0].runs[0]
        r2.font.size = Pt(9); r2.font.color.rgb = _MID
        _cell_bg(tbl.rows[i+1].cells[2], bg)

        tbl.rows[i+1].cells[3].text = "Fortnightly" if len(ids) > 1 else "Monthly"
        r3 = tbl.rows[i+1].cells[3].paragraphs[0].runs[0]
        r3.font.size = Pt(9); r3.font.color.rgb = _BODY
        _cell_bg(tbl.rows[i+1].cells[3], bg)
    _table_borders(tbl)
    for row in tbl.rows:
        _set_col_width(row.cells[0], 4.5)
        _set_col_width(row.cells[1], 4.0)
        _set_col_width(row.cells[2], 4.0)
        _set_col_width(row.cells[3], 4.0)


def _render_review_cadence(doc: Document, data: dict) -> None:
    doc.add_page_break()
    _section_heading(doc, "8.  Review Cadence & Governance")
    _body_para(doc, (
        "Risk reviews are conducted at the cadence defined below. All review outcomes are "
        "recorded in the programme audit log and retained for a minimum of 7 years in accordance "
        "with organisational document retention policy."
    ))
    _gap(doc)

    cadence = data.get("review_cadence", [])
    tbl = doc.add_table(rows=len(cadence) + 1, cols=5)
    tbl.style = "Table Grid"
    headers = ("Review Type", "Frequency", "Attendees", "Scope", "Output")
    widths  = (4.0, 2.5, 4.5, 3.5, 2.5)
    for i, (h, w) in enumerate(zip(headers, widths)):
        tbl.rows[0].cells[i].text = h
        hr = tbl.rows[0].cells[i].paragraphs[0].runs[0]
        hr.bold = True; hr.font.size = Pt(9); hr.font.color.rgb = _WHITE
        _cell_bg(tbl.rows[0].cells[i], _HDR_BG)
        _set_col_width(tbl.rows[0].cells[i], w)

    for i, entry in enumerate(cadence):
        bg = _ROW_A if i % 2 == 0 else _ROW_B
        row_vals = [
            entry.get("review_type", "—"),
            entry.get("frequency", "—"),
            entry.get("attendees", "—"),
            entry.get("scope", "—"),
            entry.get("output", "—"),
        ]
        for j, (val, w) in enumerate(zip(row_vals, widths)):
            tbl.rows[i+1].cells[j].text = val
            r = tbl.rows[i+1].cells[j].paragraphs[0].runs[0]
            r.font.size = Pt(9)
            r.font.color.rgb = _BLUE if j == 0 else (_BODY if j < 4 else _MID)
            if j == 0: r.bold = True
            _cell_bg(tbl.rows[i+1].cells[j], bg)
            _set_col_width(tbl.rows[i+1].cells[j], w)
    _table_borders(tbl)

    _gap(doc, 10)
    _body_para(doc, (
        "Any risk rated Critical must be reviewed within 24 hours of the rating change and the "
        "Programme Sponsor notified immediately via the risk escalation pathway."
    ), italic=True, color=_MUTED)


def _render_approvals(doc: Document, standards: dict) -> None:
    doc.add_page_break()
    _section_heading(doc, "9.  Approvals")
    _body_para(doc, (
        "By signing below, each approver confirms they have reviewed this Risk Registry and Risk "
        "Assessment, accept the stated risk ratings and ownership assignments, and authorise the "
        "mitigation and contingency plans documented herein."
    ))
    _gap(doc)

    approvers = [
        ("Programme Sponsor",               ""),
        ("Programme Manager",               ""),
        ("PMO Representative",              ""),
        ("Chief Risk Officer / Compliance", ""),
        ("Head of Information Security",    ""),
        ("Internal Audit Representative",   ""),
    ]
    tbl = doc.add_table(rows=len(approvers) + 1, cols=5)
    tbl.style = "Table Grid"
    for i, h in enumerate(("Role", "Name", "Signature", "Date", "Version")):
        tbl.rows[0].cells[i].text = h
        hr = tbl.rows[0].cells[i].paragraphs[0].runs[0]
        hr.bold = True; hr.font.size = Pt(9); hr.font.color.rgb = _WHITE
        _cell_bg(tbl.rows[0].cells[i], _HDR_BG)
    for row in tbl.rows:
        _set_col_width(row.cells[0], 4.5)
        _set_col_width(row.cells[1], 4.0)
        _set_col_width(row.cells[2], 4.0)
        _set_col_width(row.cells[3], 2.5)
        _set_col_width(row.cells[4], 1.5)

    for i, (role, _) in enumerate(approvers):
        bg = _ROW_A if i % 2 == 0 else _ROW_B
        tbl.rows[i+1].cells[0].text = role
        r0 = tbl.rows[i+1].cells[0].paragraphs[0].runs[0]
        r0.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = _BODY
        _cell_bg(tbl.rows[i+1].cells[0], bg)
        for ci in (1, 2, 3):
            _cell_bg(tbl.rows[i+1].cells[ci], bg)
            r = tbl.rows[i+1].cells[ci].paragraphs[0].runs[0] if tbl.rows[i+1].cells[ci].paragraphs[0].runs else tbl.rows[i+1].cells[ci].paragraphs[0].add_run()
            r.font.size = Pt(9)
        tbl.rows[i+1].cells[4].text = "v1.0"
        rv = tbl.rows[i+1].cells[4].paragraphs[0].runs[0]
        rv.font.size = Pt(9); rv.font.color.rgb = _MUTED
        _cell_bg(tbl.rows[i+1].cells[4], bg)
    _table_borders(tbl)

    _gap(doc, 10)
    _body_para(doc, (
        f"Classification: Commercial-in-Confidence  |  Retention: 7 years  |  "
        f"Owner: {standards['org']['name']} PMO Risk Team"
    ), italic=True, color=_MUTED)


# =============================================================================
# SECTION 5 — MAIN DOCX BUILDER
# =============================================================================

def build_risk_registry_docx(
    data: Dict[str, Any],
    state: PMOState,
    logo_path: Optional[str] = None,
) -> bytes:
    """
    Assemble the complete Risk Registry & Risk Assessment Word document.
    Returns raw bytes — suitable for st.download_button or HTTP response.
    """
    doc  = Document()
    _doc_defaults(doc)

    proj     = state.project
    org_name = state.standards["org"]["name"]
    risks    = data.get("risks", [])

    # Resolve logo
    if logo_path is None:
        # Look next to this file and in the config directory
        for candidate in (
            os.path.join(os.path.dirname(os.path.abspath(__file__)), "nttdata_logo.png"),
            os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.png"),
            "nttdata_logo.png",
            "logo.png",
            os.path.join("config", "logo.png"),
        ):
            if os.path.exists(candidate):
                logo_path = candidate
                break

    # ── Render all sections ───────────────────────────────────────────────────
    _cover_page(doc, proj, state.standards, logo_path)
    _render_exec_summary(doc, data, risks)
    _render_framework(doc, data)
    _render_assumptions(doc, data)
    _render_risk_summary_dashboard(doc, risks)
    _render_detailed_risks(doc, risks)
    _render_risk_mitigation_matrix(doc, risks)
    _render_risk_owners(doc, risks)
    _render_review_cadence(doc, data)
    _render_approvals(doc, state.standards)

    # ── Footer on every page ─────────────────────────────────────────────────
    _page_footer(doc, proj.project_id, org_name)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =============================================================================
# SECTION 6 — PUBLIC API  (called by nodes/generator.py)
# =============================================================================

def generate_risk_registry_artifact(state: PMOState) -> DocumentArtifact:
    """
    Full pipeline:
      1. Call LLM → structured risk JSON
      2. Build guardrail-valid markdown
      3. Build professional Word document
      4. Return populated DocumentArtifact

    The docx bytes are stashed in state.audit["risk_registry_docx"] so
    server.py /export/docx can return them without a second render.

    Usage in nodes/generator.py:
        from risk_registry_generator import generate_risk_registry_artifact
        ...
        if d == "risk_registry":
            state.docs[d] = generate_risk_registry_artifact(state)
            generated.append(d)
            continue
    """
    doc_title = state.standards["docs"]["risk_registry"].get("title", "Risk Registry")

    # Step 1 — LLM call
    try:
        risk_data = call_llm_for_risk_data(state)
        state.audit["risk_registry_llm_ok"] = True
    except Exception as e:
        state.audit["risk_registry_llm_error"] = str(e)
        # Build minimal data from project info so we always produce something
        risk_data = _validate_and_fill_risk_data({}, state.project)

    # Step 2 — Guardrail markdown
    md = build_risk_markdown(risk_data, state.project, state.standards)

    # Step 3 — Word document
    try:
        docx_bytes = build_risk_registry_docx(risk_data, state)
        state.audit["risk_registry_docx"] = docx_bytes   # available for /export/docx
    except Exception as e:
        state.audit["risk_registry_docx_error"] = str(e)
        docx_bytes = None

    art = DocumentArtifact(
        doc_type="risk_registry",
        title=doc_title,
        content_markdown=md,
        status="NOT_SUFFICIENT",   # ValidatorNode will re-validate from markdown
        reasons=["Generated by risk_registry_generator; pending validation"],
    )
    return art


# =============================================================================
# SECTION 7 — REPAIR ENTRY POINT  (called by nodes/generator.py RepairNode)
# =============================================================================

def repair_risk_registry_artifact(state: PMOState) -> DocumentArtifact:
    """
    Called by RepairNode when the ValidatorNode has flagged risk_registry as
    NOT_SUFFICIENT.  Rebuilds from a fresh LLM call with the validation issues
    appended to the prompt context.

    Usage in nodes/generator.py RepairNode.__call__:
        from risk_registry_generator import repair_risk_registry_artifact
        ...
        if d == "risk_registry":
            state.docs[d] = repair_risk_registry_artifact(state)
            repaired.append(d)
            continue
    """
    # Inject the validator findings into the audit so build_risk_llm_prompt can
    # see them via the uploaded_mapping pathway
    existing_issues = state.docs["risk_registry"].reasons if "risk_registry" in state.docs else []
    if existing_issues:
        note = "\n\nREPAIR REQUIRED. Fix these validation issues:\n" + "\n".join(
            f"  - {issue}" for issue in existing_issues
        )
        # Append to any existing uploaded text so the prompt sees it
        current = state.audit.get("uploaded_mapping", {}).get("risk_registry", "")
        state.audit.setdefault("uploaded_mapping", {})["risk_registry"] = current + note

    return generate_risk_registry_artifact(state)